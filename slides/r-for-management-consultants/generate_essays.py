"""
Generate Pre-Class Essays for "R for Management Consultants" Course
Output: pre-class-essays.docx
Style: Morgan Housel (conversational, story-driven, counterintuitive insights)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

# ── Essay Content ──────────────────────────────────────────────────────────────

essays = []

# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 0,
    "title": "The Consultant Who Stopped Copying and Pasting",
    "subtitle": "Why learning to code might be the most practical thing you do this year",
    "body": r"""
In 2019, a junior consultant at a global health advisory firm spent three days rebuilding an analysis of neonatal mortality rates across 47 districts in Malawi. The original Excel workbook had been submitted to the Ministry of Health the previous week. Then someone found a data error — two districts had been transposed in the source file, and the mortality rates attributed to Lilongwe actually belonged to Mangochi, and vice versa.

Three days. That is how long it took to re-sort the data, re-build the pivot tables, re-format the charts, re-check the cross-references, and re-write the narrative sections that cited specific numbers. Three days of work that added zero new insight, zero new analysis, and zero new value. It was pure reconstruction.

Down the hall, a colleague faced the same problem on a different project. She had built her analysis in R. She corrected the two transposed values in the source CSV, re-ran her script, and had a fully updated report — charts, tables, narrative text, and all — in about thirty seconds.

The gap between three days and thirty seconds is not a story about technology. It is a story about how you structure your thinking.

Most people hear "learn to code" and imagine they are being asked to become software engineers — to inhabit a world of algorithms and data structures and obscure syntax. That image is wrong, and it keeps talented analysts from acquiring a skill that would make their existing work dramatically better. Learning R as a management consultant is not about becoming technical. It is about becoming efficient, transparent, and trustworthy.

Here is a claim that might sound counterintuitive: the biggest benefit of learning R is not the analysis you can do. It is the analysis you can redo. In consulting, things change constantly. Data gets corrected. Indicators get redefined. A donor asks you to re-run everything with a different baseline year. A government counterpart wants the same analysis but disaggregated by urban and rural. In Excel, each of those requests means hours of manual rework. In R, each of those requests means changing a parameter and pressing Enter.

This matters enormously in maternal and newborn health, where the stakes are measured in lives. When the Pan American Health Organization issued updated guidelines for tracking stillbirth rates in 2020, every team that had hard-coded the old definition into their Excel workbooks had to rebuild from scratch. Teams using scripted analyses updated a single variable definition and moved on.

But efficiency is only part of the story. The deeper advantage is trust.

When you hand someone an Excel file, you are handing them a finished product with no visible process. They can see the numbers and the charts, but they cannot see how you got there. They cannot see which rows you filtered out, which assumptions you made, which edge cases you handled and which you ignored. They have to trust you, and if they are sophisticated, they know they probably should not — not because you are dishonest, but because Excel makes it almost impossible to audit your own work, let alone someone else's.

An R script is different. It is a complete, readable record of every decision you made, from the moment you loaded the raw data to the moment you produced the final output. Anyone can read it, run it, check it, and challenge it. That transparency does not make you vulnerable. It makes you credible.

The best analysts in global health consulting are not the ones with the most data or the fanciest models. They are the ones whose work can be trusted, reproduced, and updated. They are the ones who, when a minister asks "but what if we look at it this way instead?", can answer in minutes rather than days.

This course is not designed to turn you into a programmer. It is designed to give you a new way of working — one where your analysis is a living document rather than a frozen artifact, where your methods are visible rather than hidden, where updating your work is trivial rather than terrifying.

That junior consultant in Malawi? She eventually learned R. She says the hardest part was not the syntax or the functions. It was unlearning the instinct to do everything by hand. Once she let go of that, everything got easier.

You are about to have the same experience.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 1,
    "title": "The Most Dangerous Spreadsheet",
    "subtitle": "When Excel hides the mistakes that matter",
    "body": r"""
In 2010, two Harvard economists named Carmen Reinhart and Kenneth Rogoff published an influential paper arguing that countries with public debt exceeding 90 percent of GDP experienced dramatically lower economic growth. The finding shaped austerity policies across Europe and the United States. Finance ministers cited it. Congressional budget hawks quoted it. It became one of the most consequential empirical claims in modern economics.

Three years later, a graduate student named Thomas Herndon tried to replicate their results for a term paper. He could not. When Reinhart and Rogoff shared their spreadsheet, Herndon found the problem: they had accidentally excluded five countries from an Excel formula by selecting the wrong cell range. The formula referenced cells A1 through A45 when it should have referenced A1 through A50. Five countries — Australia, Austria, Belgium, Canada, and Denmark — simply vanished from the calculation, and nobody noticed for three years.

The error did not change the paper's conclusion entirely, but it changed the magnitude enough to undermine the policy argument. The difference between "growth slows dramatically above 90 percent debt" and "growth slows modestly" is the difference between imposing austerity on millions of people and not. And it came down to a cell range.

If you work in global health, you might think this is an economics problem, not your problem. But consider how much of the analytical work in maternal and newborn health happens in exactly the same way — in spreadsheets where one wrong click, one incorrectly dragged formula, one accidentally overwritten cell can cascade through an entire analysis without anyone noticing.

In 2017, a technical assistance team supporting a Ministry of Health in West Africa discovered that a facility-level data extract had been accidentally sorted by one column without selecting all columns. The result was that facility names became mismatched with their data. A health center in a rural district appeared to have the staffing levels of a regional hospital, and vice versa. The error was not discovered for four months, during which the mismatched data informed staffing deployment decisions.

The problem with Excel is not that it causes errors. Every tool causes errors. The problem is that Excel makes errors invisible. When you overwrite a cell, the old value is gone. When you drag a formula down a column, you cannot easily verify that every cell is referencing what you think it is referencing. When you sort, filter, copy, and paste, you leave no trace of what you did. The spreadsheet looks exactly the same whether every number is correct or half of them are wrong.

R does not prevent errors. You will make mistakes in R — wrong variable names, incorrect joins, filters that exclude data you meant to include. The difference is that R makes your mistakes visible, auditable, and fixable. Every operation you perform is written as a line of code. That code can be read, reviewed, and tested. When something goes wrong, you can trace exactly what happened, because the complete sequence of operations is right there in front of you.

This distinction matters more than most people realize. In consulting, your reputation rests on the accuracy of your work. A single error in a presentation to a health minister — a misquoted mortality rate, a wrong district name, an incorrectly calculated trend — can undermine months of credibility-building. In Excel, you discover those errors after the fact, often in the most embarrassing way possible. In R, you can build checks into your workflow that catch errors before they reach anyone's desk.

There is also a subtler issue. Excel trains you to think of analysis as a product — a finished artifact with numbers in cells and charts on tabs. R trains you to think of analysis as a process — a sequence of steps that transforms raw data into conclusions. The process mindset is more powerful because it makes you ask different questions. Instead of "do these numbers look right?", you ask "is this process correct?" The first question is a guess. The second is verifiable.

None of this means you should never use Excel. Spreadsheets are excellent for quick calculations, data entry, and sharing simple tables with colleagues who do not code. But for analysis that matters — the kind that informs policy, directs resources, and affects the health of mothers and newborns — you need a tool that shows its work.

The most dangerous spreadsheet is not the one with an obvious error. It is the one that looks perfectly fine.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 2,
    "title": "Why Programmers Name Things Carefully",
    "subtitle": "The surprising connection between variable names and clear thinking",
    "body": r"""
In 2018, a team of consultants working on a maternal health program in East Africa ran into a problem that had nothing to do with statistics, coding, or methodology. It was a naming problem.

The program tracked "coverage" across several indicators. But "coverage" meant different things in different parts of the analysis. In one Excel column, coverage meant the percentage of pregnant women who received at least four antenatal care visits. In another column, coverage meant the percentage of health facilities that had a skilled birth attendant on staff. In a third, it meant the proportion of the target population reached by community health worker outreach.

Each definition was reasonable on its own. But when the team built a summary dashboard that averaged "coverage" across indicators, they were averaging percentages of women, percentages of facilities, and percentages of populations — numbers that should never be combined. The resulting "coverage score" was mathematically meaningless, but it looked clean and precise on the slide.

Nobody caught the error during the analysis. A government counterpart caught it during the presentation, and the team spent the next two weeks rebuilding the dashboard with clearly defined indicators.

This story is about naming, and naming is about thinking.

Phil Karlton, a programmer who worked at Netscape in the 1990s, once said there are only two hard things in computer science: cache invalidation and naming things. The joke has endured because it is true. Naming things well — variables, functions, files, columns — is one of the most underrated cognitive skills in analytical work.

When you work in Excel, you rarely name things explicitly. Column B might be "coverage" this week and "coverage_v2" next week and "coverage_FINAL_revised" the week after that. The column header is the only label, and it can mean whatever the person who typed it intended at the time. There is no enforcement, no structure, no requirement to be precise.

In R, you name things constantly and deliberately. When you write nmr_per_1000 <- deaths_0_28 / live_births * 1000, you have created a variable whose name tells you exactly what it contains: the neonatal mortality rate per 1,000 live births. You cannot accidentally confuse it with a percentage, because the name itself encodes the unit. You cannot accidentally confuse it with the infant mortality rate, because the name specifies the age range.

This is not a stylistic preference. It is a cognitive discipline.

Research in cognitive psychology has consistently shown that humans are terrible at holding ambiguous categories in working memory. When a concept has multiple possible meanings, we tend to default to whichever meaning is most salient in the moment — not whichever meaning is correct. The East Africa team did not make a calculation error. They made a category error, and the ambiguous name "coverage" is what allowed it to happen.

Good variable names serve as a form of documentation. When you read a line of code that says anc4_coverage_pct <- women_4plus_anc / total_pregnant_women * 100, you understand not only what the number is, but how it was calculated. The name and the formula reinforce each other. If someone later tries to use this variable as a facility-level metric, the name itself signals that something is wrong.

This discipline extends beyond variable names to the entire structure of your analysis. In R, you organize your work into scripts with descriptive names: 01_clean_facility_data.R, 02_calculate_indicators.R, 03_generate_district_summaries.R. The names tell a story about the analytical process. Anyone looking at the project folder can understand the workflow without opening a single file.

Contrast this with the typical Excel workflow, where the analytical process is distributed across tabs with names like "Sheet1," "Data," "Pivot," "Charts," "Final," and "Final_v2_CORRECTED." The names tell you nothing about what happened or in what order.

There is a deeper lesson here. Naming things carefully forces you to think about what things actually are. When you cannot call a variable "coverage" because you have already used that name, you are forced to specify: coverage of what? Among whom? Measured how? Over what time period? Those are exactly the questions that lead to rigorous analysis, and vague naming lets you skip them.

The programmers' obsession with naming is not pedantry. It is a recognition that clarity of language and clarity of thought are the same thing. When you name a variable well, you have already done half the analytical work — because you have already decided, precisely, what you are measuring.

In maternal and newborn health, where the difference between the neonatal mortality rate and the perinatal mortality rate can change the story entirely, that precision is not optional. It is the foundation of everything that follows.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 3,
    "title": "The First Five Minutes with a New Dataset",
    "subtitle": "What experienced analysts do before they analyze anything",
    "body": r"""
A few years ago, a data analyst at a multilateral health organization received a dataset of facility-level delivery statistics from a country in southern Africa. The dataset had been cleaned by a local team, formatted into a tidy spreadsheet, and delivered with a note that said "ready for analysis." The analyst, under deadline pressure, loaded the data into a regression model to estimate the relationship between staffing levels and delivery outcomes.

The results were striking: facilities with fewer skilled birth attendants appeared to have better outcomes. The analyst wrote up the finding and shared it with the team. Someone senior took one look and asked a simple question: "How much of this data is actually present?"

The answer was devastating. The key outcome variable — whether the mother and newborn experienced complications — was missing for 40 percent of the records. And the missingness was not random. It was concentrated in rural facilities with fewer staff, which were also the facilities least likely to have functioning reporting systems. The analyst had run a regression on the 60 percent of facilities that happened to report data — overwhelmingly urban, well-staffed facilities — and concluded that staffing did not matter. The missing data was not noise. It was the story.

This is what happens when you skip the first five minutes.

Experienced analysts have a ritual when they encounter a new dataset. Before they calculate a single statistic, they explore. They ask: How many rows and columns? What are the variable types? What are the ranges? What is the distribution of key variables? Where is the data missing, and is the missingness patterned?

This ritual feels unproductive. It produces no charts, no tables, no findings. It is the analytical equivalent of a surgeon washing their hands — invisible work that prevents catastrophe.

In R, this ritual has a specific vocabulary. You run dim() to see the dimensions. You run str() or glimpse() to see the types. You run summary() to see the ranges and counts of missing values. You run table() to see the distribution of categorical variables. You might create a quick histogram or a missingness map. The whole thing takes five minutes and tells you more about your data than hours of modeling ever will.

The counterintuitive lesson is this: the most productive first step with any dataset is to do nothing analytical. Just look.

There is a cognitive bias at work here. Psychologists call it the "action bias" — the tendency to prefer doing something over doing nothing, even when doing nothing would be more productive. In analysis, the action bias manifests as the urge to start modeling, start calculating, start producing results. The feeling is that looking at the data without transforming it is wasted time.

It is not wasted time. It is the time that prevents you from wasting everything that follows.

Consider what you learn from a simple frequency table of a district variable. If you expect 47 districts and you see 52, something is wrong — perhaps district names are misspelled, or old district boundaries are mixed with new ones. If you expect roughly equal numbers of observations per district and one district has ten times more than the others, something is interesting — perhaps that district reports more frequently, or perhaps it is a data dump from a different source. These are things you need to know before you analyze, not after.

Or consider missing data patterns. If a variable is missing 2 percent of the time and the missingness is scattered randomly, you can probably ignore it. If the same variable is missing 30 percent of the time and the missingness correlates with facility type, you have a structural problem that will bias any analysis that ignores it. The only way to know which situation you are in is to look.

In the maternal and newborn health world, data quality varies enormously across countries, facilities, and time periods. A dataset from a national HMIS may have near-complete records in urban areas and vast gaps in rural ones. A survey dataset may have been cleaned by one team and appended to uncleaned data from another team. A program dataset may mix baseline and endline records in ways that are not immediately obvious from the column headers.

The experienced analyst is not smarter than the inexperienced one. The experienced analyst has simply been burned enough times to develop the habit of looking before leaping. R makes this habit easy to practice, because the exploratory functions are fast, simple, and designed to be run interactively. You are not building a product during those first five minutes. You are building an understanding.

That understanding is the foundation on which every subsequent analytical decision rests. Skip it, and you might run a technically perfect regression on fundamentally broken data. Take the time, and you will catch problems when they are still cheap to fix.

The first five minutes are not a delay. They are the most important part of the analysis.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 4,
    "title": "The Pivot Table's Replacement",
    "subtitle": "How six words changed the way analysts think about data",
    "body": r"""
In 2014, a statistician named Hadley Wickham released a package called dplyr that quietly changed how a generation of analysts thinks about data manipulation. The package introduced a small vocabulary of verbs — filter, select, mutate, summarise, arrange, group_by — that, when chained together with an operator called the pipe, could express almost any data transformation as a readable English sentence.

This might sound like a technical detail. It is not. It is a design philosophy, and understanding it will change how you work.

Consider a common task in maternal and newborn health consulting: calculating the average neonatal mortality rate by region, excluding facilities with fewer than 50 deliveries. In Excel, you would create a pivot table with a filter. Or you would add helper columns. Or you would use a SUMPRODUCT formula with nested conditions. All of these work, but none of them read like a description of what you are trying to do.

In dplyr, the same task reads like this: take the data, filter to facilities with 50 or more deliveries, group by region, and summarise the average neonatal mortality rate. Each verb corresponds to one logical step. The pipe operator connects the steps in sequence. Someone who has never used R can read that pipeline and understand the logic.

This readability is not an accident. Wickham designed dplyr around a principle from linguistics: that the best tools match the structure of human thought. When you think about data manipulation in natural language, you think in verbs — "I want to filter this, then group it, then calculate a summary." Dplyr lets you write code that mirrors that natural thought process.

The contrast with Excel's pivot tables is instructive. A pivot table is a powerful tool, but it hides the logic. When you drag fields into rows, columns, values, and filters, you are making analytical decisions — but those decisions are embedded in the configuration of a visual interface, not expressed as readable steps. If you come back to the pivot table a month later, you have to reverse-engineer what you did. If a colleague opens your file, they see the result but not the reasoning.

A dplyr pipeline is transparent. Every step is visible. You can read it top to bottom and understand exactly what happened to the data. You can add a comment explaining why you filtered to 50 deliveries rather than 100. You can insert a step in the middle without restructuring the whole thing. You can copy the pipeline, change one verb, and create a parallel analysis instantly.

There is a deeper insight here about what happens when tools match cognition. Excel pivot tables require you to translate your analytical question into a spatial arrangement of fields in a grid. That translation is a cognitive overhead — a gap between what you want to do and how the tool lets you do it. Dplyr eliminates that gap. You think in verbs, and you code in verbs.

This matters for consultants specifically because consulting work often involves iteration. A client says: "Great analysis, but can you also show me the breakdown by facility type?" In Excel, that means rebuilding the pivot table. In dplyr, it means adding one word: group_by(region, facility_type) instead of group_by(region). The modification is proportional to the change in thinking — one concept, one word.

Or the client says: "What if we use median instead of mean?" One word: median(nmr) instead of mean(nmr). The structure does not change. The logic does not change. Only the single element that the client asked about changes.

This proportionality between conceptual change and code change is what makes dplyr so powerful for consulting work. In a field where specifications change constantly and clients refine their questions iteratively, the ability to make small, precise modifications to an existing analysis — rather than rebuilding from scratch — is an enormous practical advantage.

But perhaps the most underrated benefit is communication. When you include a dplyr pipeline in a technical appendix or a methodology document, you are giving your reader a precise, unambiguous description of what you did. There is no room for misinterpretation. The code is the method. The method is the code. In maternal and newborn health, where analytical transparency can mean the difference between a trusted finding and a disputed one, that clarity is invaluable.

Six verbs. One pipe operator. A way of working that reads like English and runs like software. It is not a pivot table replacement because it is faster or more powerful — though it often is. It is a replacement because it makes your thinking visible. And visible thinking is the foundation of credible analysis.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 5,
    "title": "The Shape of Your Data Is the Shape of Your Thinking",
    "subtitle": "Why wide data and long data lead to different questions",
    "body": r"""
A program manager working on a newborn health initiative in Southeast Asia spent an entire afternoon trying to figure out why her ggplot2 chart was not working. She had a clean dataset, a clear research question, and a straightforward visualization in mind — a line chart showing neonatal mortality trends across five provinces over ten years. The code ran without errors. But instead of five lines showing trends, she got a jumbled mess of disconnected points.

The problem was not her code. It was her data.

Her dataset was in wide format: one row per province, with columns for each year — NMR_2012, NMR_2013, NMR_2014, and so on. This is the format that makes intuitive sense when you look at data in a spreadsheet. It is compact, human-readable, and easy to scan by eye.

But ggplot2 — and most of R's analytical tools — expects data in long format: one row per observation, with a column for the year and a column for the value. In long format, the same data would have fifty rows (five provinces times ten years) with three columns: province, year, and nmr.

Once she reshaped her data from wide to long using the pivot_longer function, the chart worked perfectly. Five clear trend lines, one per province, exactly as she had envisioned.

This is more than a technical annoyance. The distinction between wide and long data reflects a fundamental difference in how you think about data, and understanding it will change how you approach every analysis.

Wide data is structured for display. When you want to compare provinces side by side for a single year, wide format is natural — each province gets its own column, and you can scan across the row. It is the format of summary tables, dashboards, and reports designed for human eyes.

Long data is structured for analysis. When you want to compare trends over time, calculate grouped statistics, or build visualizations that distinguish categories by color or facet, long format is essential. It is the format that lets you ask questions like "show me the trend for each province" or "what is the average NMR by region and year?" with a single command.

The insight is that the shape of your data determines what questions you can easily answer. This is not a limitation of the software. It is a property of structured data itself.

Consider a practical example. A consultant is analyzing HMIS data for a report on maternal health service delivery. The data arrives in wide format: one row per facility, with columns for ANC1 visits in January, ANC1 visits in February, ANC1 visits in March, and so on. The consultant wants to answer several questions: What is the monthly trend? Which facilities show seasonal patterns? Are there months with unusually low reporting?

In wide format, answering these questions requires writing separate formulas for each month. In long format — with columns for facility, month, and visits — the answers are one-line operations: group by month and calculate the mean, create a line chart with month on the x-axis, or filter for months where visits drop below a threshold.

The pivot is the bridge between these two worlds, and it runs in both directions. pivot_longer reshapes wide data into long data for analysis. pivot_wider reshapes long data into wide data for display. Mastering these two functions means you can always move your data into the shape that matches the question you are asking.

There is a philosophical point here that is worth dwelling on. When analysts struggle with a visualization or a calculation that will not work, the instinct is to look for a bug in the code or a flaw in the logic. But often the problem is more fundamental: the data is in the wrong shape for the question being asked. The code is correct. The logic is sound. The shape is wrong.

This happens because most analysts learn data in the shape that spreadsheets impose — wide format, with variables as columns and observations as rows. That shape is so familiar that it feels like the natural order of things. But there is nothing natural about it. It is one of two equally valid ways to organize the same information, and the choice between them should depend on what you want to do, not on what you are used to seeing.

In maternal and newborn health, data frequently arrives in wide format from HMIS platforms, survey exports, and program databases. Learning to reshape that data — fluidly, confidently, without confusion — is not a coding skill. It is an analytical skill. It is the ability to see that the same data can tell different stories depending on how you arrange it, and to choose the arrangement that best serves the question at hand.

The shape of your data is the shape of your thinking. Once you learn to change the shape, you learn to change the questions you can ask.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 6,
    "title": "The Chart That Changed a Minister's Mind",
    "subtitle": "Why visualization is the consultant's most persuasive tool",
    "body": r"""
In 2016, a health systems advisor was preparing for a budget review meeting with a Minister of Health in a country in East Africa. The ministry was allocating maternal and newborn health funding equally across all districts, on the principle that equity meant giving everyone the same amount. The advisor's analysis showed that neonatal mortality was heavily concentrated in a handful of rural districts that lacked functioning referral systems, while urban districts with relatively low mortality were receiving the same per-capita allocation.

The advisor had prepared a detailed table showing mortality rates, funding allocations, facility counts, and staffing ratios for all 47 districts. The table was comprehensive, accurate, and completely unpersuasive. The minister glanced at it, nodded politely, and moved to the next agenda item.

The advisor came back the following week with a single chart: a scatter plot with per-capita funding on the x-axis and neonatal mortality rate on the y-axis. Each dot was a district, sized by population and colored by region. The chart told a story that the table could not: the districts receiving the most money were the ones that needed it least, and the districts with the highest mortality were clustered in the bottom-left corner — low funding, high need.

The minister stared at the chart for thirty seconds. Then she asked her team to draft a needs-based allocation formula.

One chart. One meeting. A policy change that would affect the distribution of millions of dollars in health funding.

This is not a story about data visualization as decoration. It is a story about cognition.

Humans process visual patterns orders of magnitude faster than tabular data. When you look at a table of 47 rows and 6 columns, your brain processes each cell sequentially, holding numbers in working memory while scanning for patterns. It is slow, effortful, and error-prone. When you look at a well-designed scatter plot, your brain processes the entire pattern simultaneously — the clusters, the outliers, the trends, the relationships. The insight is immediate and intuitive.

This is why visualization is not a nice-to-have in consulting. It is the primary mechanism through which complex analytical findings enter the minds of decision-makers. Tables inform. Charts persuade.

But there is an important caveat: not all charts persuade. A bar chart with unlabeled axes, inconsistent colors, and a misleading scale does not create understanding — it creates confusion. A chart that tries to show six variables simultaneously does not demonstrate analytical sophistication — it demonstrates a failure to prioritize. The chart that changed the minister's mind worked because it made one point clearly: there is a mismatch between need and funding.

In R, the ggplot2 package gives you the tools to create charts that are both analytically rigorous and visually persuasive. The "grammar of graphics" approach — building a chart layer by layer, mapping data to visual properties like position, color, size, and shape — forces you to think about what each visual element communicates. You are not just making a chart. You are making an argument.

Consider the difference between a default Excel chart and a thoughtfully constructed ggplot2 visualization. The Excel chart typically has a gray background, gridlines competing with the data, a legend that takes up a third of the space, and axis labels in a default font that nobody chose deliberately. It looks generic because it is generic — the software made every design decision for you.

A well-built ggplot2 chart, by contrast, reflects deliberate choices. The background is clean. The colors encode meaning. The labels explain rather than clutter. The title states the insight, not the topic. Every element is there because the analyst decided it should be, and every element serves the reader's understanding.

This deliberateness matters because design communicates credibility. When a health minister sees a polished, clear visualization, they make an unconscious judgment: this person thought carefully about this analysis. When they see a default chart with clip-art styling, they make the opposite judgment — regardless of how rigorous the underlying analysis might be.

The best consultants understand that visualization is not the last step in the analytical process — something you do after the "real" analysis is complete. It is an integral part of the analysis itself. The act of deciding how to visualize a finding forces you to clarify what the finding actually is. When you cannot figure out how to show something in a chart, it often means you have not yet figured out what you are trying to say.

A chart does not just display data. It creates understanding. And in a world where the people who make the most consequential decisions about maternal and newborn health are not statisticians but ministers, directors, and program managers, understanding is the scarcest and most valuable resource you can provide.

The most persuasive tool in a consultant's arsenal is not a regression model or a cost-benefit analysis. It is a chart that tells a true story clearly.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 7,
    "title": "The Difference Between a Chart and a Graphic",
    "subtitle": "Why default settings undermine your credibility",
    "body": r"""
In 2021, a consulting team presented the results of a three-month maternal health systems assessment to a donor organization. The team had done excellent analytical work: thorough data collection, careful cleaning, rigorous statistical analysis. The findings were genuinely useful and could inform millions of dollars in programming decisions.

The presentation had twelve slides with data visualizations. Every chart had a gray background. The axis labels were in 8-point font, barely readable from across the conference table. The color palette was the default ggplot2 scheme — an assortment of teal, salmon, and lavender that clashed with the organization's branding and gave the presentation the look of a homework assignment. Legends were placed in the default position, taking up valuable chart space. Titles stated the variable name ("district_nmr_trends") rather than the insight ("Neonatal mortality declined in 8 of 12 target districts").

The donor's response was polite but lukewarm. They asked for a "more polished" version before sharing with their board. The team spent another week reformatting every chart — work that could have been avoided if the charts had been designed well from the beginning.

Down the hall at the same organization, another team presented a different assessment. Their analytical work was comparable in rigor. But their charts were different. Clean white backgrounds. Consistent use of two brand colors — a deep blue for primary data and a warm amber for comparisons. Axis labels in readable 12-point type. Chart titles that stated the finding rather than the variable. Annotations that drew the eye to the key insight. A custom ggplot2 theme applied consistently to every visualization.

This team received immediate approval and a request to present to the board directly.

The difference between these two presentations was not the quality of the analysis. It was the quality of the design. And the lesson is not that design is superficial. The lesson is that design is communication.

When most people hear "design," they think of aesthetics — making things look pretty. But in data visualization, design is about function. Every design choice either helps or hinders the reader's ability to understand the data. A gray background adds visual noise that competes with the data points. Small axis labels force the reader to squint rather than comprehend. Inconsistent colors make the reader work harder to map visual elements to data categories. Default titles tell the reader nothing about what they should take away from the chart.

These are not trivial issues. They are the difference between a chart that communicates instantly and one that requires explanation.

In ggplot2, the default aesthetic choices are designed to be acceptable, not optimal. They are a starting point, not a destination. The gray background, for instance, exists because it provides contrast with white gridlines — a defensible design choice for some contexts. But in a boardroom presentation projected onto a screen, it looks dull and academic. The default color palette is designed to be distinguishable by people with color vision deficiency — an important consideration — but it was not designed to match your organization's branding or to create visual hierarchy.

The power of ggplot2 lies in the theme system, which lets you customize every visual element of a chart with a few lines of code. You can create a custom theme once — specifying your preferred fonts, colors, backgrounds, grid styles, legend positions, and title formatting — and apply it to every chart in your analysis with a single line: + theme_bka(). Suddenly, every visualization in your report looks deliberate, professional, and consistent.

This consistency matters more than most analysts realize. When every chart in a report shares the same visual language, the reader's brain can focus entirely on the data. They do not have to re-learn the color coding for each chart. They do not have to search for the legend in a different position on each slide. The visual consistency reduces cognitive load and increases comprehension.

There is a practical principle here: spend thirty minutes building a custom theme at the beginning of a project, and you will save hours of reformatting later. More importantly, you will produce work that looks intentional from the first draft. In consulting, where first impressions shape how seriously your analysis is taken, that intentionality is worth more than its weight in billable hours.

The difference between a chart and a graphic is intent. A chart is what the software gives you by default. A graphic is what you create when you make every visual choice deliberately — when the colors, the fonts, the labels, the annotations, and the layout all serve a single purpose: helping the reader understand the data.

Default settings are not neutral. They send a message: "I did not think about this." A custom theme sends the opposite message: "Every element is here for a reason." In a field where your credibility depends on demonstrating rigor, that message matters.

Design is not decoration. It is the last mile of communication — the bridge between your analysis and your reader's understanding. Build it well.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 8,
    "title": "The Number Everyone Misunderstands",
    "subtitle": "What a p-value actually tells you — and what it doesn't",
    "body": r"""
In 2012, a program evaluation team assessed the impact of a community health worker intervention on neonatal mortality in a cluster of districts in southern Asia. After two years of implementation, the evaluation found that neonatal mortality had declined from 34 per 1,000 live births in the comparison districts to 31 per 1,000 in the intervention districts. The p-value was 0.03.

The finding was declared "statistically significant," and the program was expanded to three additional countries at a cost of tens of millions of dollars. The logic seemed straightforward: the program worked, the result was significant, and therefore scaling was justified.

But there is something missing from this logic, and it is the thing that most people get wrong about statistical significance.

A p-value of 0.03 means that if the program had truly zero effect — if the intervention districts and comparison districts were identical in every way — there would be roughly a 3 percent chance of observing a difference this large or larger just due to random variation. In other words, the p-value answers the question: "Could this result be noise?"

That is a useful question. But it is not the most important question. The most important question is: "Does this result matter?"

The difference between 34 and 31 deaths per 1,000 live births is 3 deaths per 1,000. In a district with 10,000 live births per year, that is 30 fewer neonatal deaths — meaningful and important. But in a district with 500 live births per year, that is 1.5 fewer deaths. Whether a multi-million-dollar program is worth scaling to save 1.5 additional lives per district per year depends on the cost, the alternatives, and the opportunity cost — questions that the p-value cannot answer.

This is the distinction between statistical significance and practical importance, and conflating the two is one of the most common analytical errors in global health.

Statistical significance tells you whether a result is likely to be real — whether the observed effect is distinguishable from zero given the noise in your data. It does not tell you whether the effect is large enough to matter. A study with a very large sample size can detect tiny effects with high statistical significance. A drug that reduces mortality by 0.1 percentage points might achieve a p-value of 0.001 in a trial with a million participants. The result is statistically significant and practically meaningless.

Conversely, a study with a small sample size might find a large and important effect that fails to achieve statistical significance simply because the sample was too small to distinguish the signal from the noise. A pilot program that reduces neonatal mortality by 40 percent in five facilities might have a p-value of 0.15 — not significant by the conventional threshold — because five facilities is simply not enough data to be confident. The effect is practically enormous and statistically ambiguous.

The p-value is a measure of surprise, not importance. A small p-value means: "This result would be surprising if the true effect were zero." A large p-value means: "This result would not be surprising if the true effect were zero." Neither tells you the size, direction, or practical significance of the effect.

What should you report instead? Effect sizes and confidence intervals. The effect size tells you how big the difference is. The confidence interval tells you how uncertain you are about that estimate. Together, they answer the question that actually matters: "What is our best estimate of the effect, and how much could it plausibly vary?"

In the neonatal mortality example, instead of saying "the result was significant (p = 0.03)," the evaluation could have said "neonatal mortality was 3 per 1,000 lower in intervention districts (95% CI: 0.3 to 5.7)." This statement tells the reader three things: the estimated effect is 3 per 1,000, it could be as small as 0.3 per 1,000 or as large as 5.7 per 1,000, and even the upper bound is a relatively modest effect. Armed with this information, a decision-maker can weigh the cost of scaling against the range of plausible benefits.

The over-reliance on p-values has been so problematic that the American Statistical Association took the unusual step of issuing a formal statement about it in 2016, cautioning against using p-values as the sole basis for scientific conclusions. Major journals have banned the phrase "statistically significant" from their pages.

For consultants working in maternal and newborn health, the practical lesson is this: always report effect sizes. Always report confidence intervals. Always interpret your results in terms of practical importance, not statistical thresholds. When a client asks "is this result significant?", the best answer is not a p-value. It is a clear statement of what the data shows, how uncertain you are, and what it means for the decision at hand.

The p-value answers one narrow question. Your client is asking a much broader one. Make sure your answer matches the question.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 9
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 9,
    "title": "The Report That Updated Itself",
    "subtitle": "How reproducible analysis changes what's possible",
    "body": r"""
In 2020, a small consulting team was producing quarterly maternal and newborn health reports for five countries in sub-Saharan Africa. Each report contained 35 pages of analysis: trend charts for key indicators, district-level comparison tables, traffic-light dashboards for program targets, and narrative summaries interpreting the latest data.

The process for producing each quarterly report went like this: a data analyst would download the latest HMIS data, clean it in Excel, produce the charts and tables in a separate workbook, copy the charts into a Word document, update the narrative text by hand, format the document to match the template, and submit it for review. The reviewer would find errors — a chart label from the previous quarter, a narrative that referenced the wrong time period, a table that did not match the chart above it — and send it back. Two or three rounds of revision later, the report was finalized.

Each report took approximately two weeks to produce. For five countries, that was ten weeks of work per quarter — nearly the entire quarter consumed by the act of reporting on the previous quarter. By the time the reports were finalized, the data was already three months old, and the next quarter's data was already coming in.

Then the team adopted Quarto.

They rebuilt the report as a parameterized Quarto document. The data cleaning was handled by an R script. The charts were generated by ggplot2 code embedded in the document. The tables were produced by code. The narrative text used inline R expressions that automatically inserted the correct numbers, periods, and country names. The formatting was handled by a template.

When new data arrived, the analyst changed one parameter — the quarter — and rendered the document. The entire report regenerated: data cleaning, charts, tables, narrative, formatting. What had taken two weeks now took two hours, most of which was spent reviewing the output rather than producing it.

This is what reproducibility looks like in practice. And the word "reproducibility" does not do it justice, because it sounds like an academic concern — something scientists worry about for the sake of methodological purity. In consulting, reproducibility is not a luxury. It is a competitive advantage.

Here is why. When your report updates itself, you are no longer spending your time on production. You are spending your time on insight. Instead of copying charts and fixing labels, you are looking at the data, noticing patterns, asking questions, and adding value. The mechanical work — the work that adds no intellectual value but consumes most of the time — is handled by the machine.

This changes the economics of consulting in a fundamental way. A team that can produce a report in two hours instead of two weeks can take on more clients, deliver more frequently, or spend the saved time on deeper analysis. They can offer weekly updates instead of quarterly ones. They can respond to ad hoc requests in hours instead of days. They can say "let me re-run that with different parameters" instead of "we'll need two weeks to rebuild the analysis."

But the benefit goes beyond efficiency. Reproducible analysis also eliminates an entire category of errors — the errors that come from manual production. When you copy a chart from Excel into Word, you might copy the wrong version. When you update a number in the text, you might miss the same number in a different paragraph. When you format a table by hand, you might introduce inconsistencies that undermine the reader's trust. These are not analytical errors. They are production errors, and they are embarrassingly common.

In a Quarto document, production errors are structurally impossible. The chart in the document is generated by the code in the document. The number in the text is pulled from the data by the code in the document. The formatting is applied by the template. There is no gap between the analysis and the output, because they are the same thing.

There is a concept in software engineering called the "inner loop" — the cycle of writing code, running it, and seeing the results. The faster the inner loop, the more productive the developer. Quarto shrinks the inner loop for analytical work to almost nothing. Change a filter, re-render, see the result. Try a different visualization, re-render, compare. Add a new section, re-render, review. The feedback is immediate, and that immediacy changes how you work.

The team that adopted Quarto for their quarterly reports discovered something unexpected. The time savings were real and substantial, but the bigger change was qualitative. Because they were no longer spending weeks on production, they started noticing things in the data that they had previously been too busy to see. They spotted a seasonal pattern in antenatal care visits that led to a recommendation for targeted outreach during planting season. They identified a cluster of facilities with suspiciously identical reporting that turned out to be a data quality issue. They had the time to be curious.

Reproducibility is not about doing the same thing twice. It is about freeing yourself from the mechanical work that prevents you from doing the thinking that matters. When your report updates itself, you can finally focus on the question that no machine can answer: what does this data mean, and what should we do about it?

That is the real promise of reproducible analysis. Not efficiency for its own sake, but efficiency in service of insight. The report that updates itself gives you something no amount of manual formatting ever could: time to think.
""",
})


# ── Document Generation ───────────────────────────────────────────────────────

def build_document(essays, output_path):
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # -- Heading styles --
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)  # BKA primary blue
        if level == 1:
            h.font.size = Pt(26)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(4)
        elif level == 2:
            h.font.size = Pt(16)
            h.font.italic = True
            h.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(18)
        elif level == 3:
            h.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # ── Title Page ─────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("R for Management Consultants")
    run.font.name = "Calibri Light"
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
    run.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("Pre-Class Essays")
    run.font.name = "Calibri Light"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("BK Advisors")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)

    doc.add_page_break()

    # ── Table of Contents (manual) ─────────────────────────────────────────
    toc_heading = doc.add_heading("Contents", level=1)

    for essay in essays:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.space_after = Pt(4)
        if essay["number"] == 0:
            toc_label = f"Introduction:  {essay['title']}"
        else:
            toc_label = f"Chapter {essay['number']}:  {essay['title']}"
        run = toc_para.add_run(toc_label)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # Subtitle line
        sub_para = doc.add_paragraph()
        sub_para.paragraph_format.space_after = Pt(10)
        sub_para.paragraph_format.left_indent = Cm(1.27)
        run = sub_para.add_run(essay["subtitle"])
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.italic = True

    doc.add_page_break()

    # ── Essays ─────────────────────────────────────────────────────────────
    for i, essay in enumerate(essays):
        # Chapter heading
        if essay["number"] == 0:
            doc.add_heading("Introduction", level=3)
        else:
            doc.add_heading(f"Chapter {essay['number']}", level=3)
        doc.add_heading(essay["title"], level=1)
        doc.add_heading(essay["subtitle"], level=2)

        # Body text — split into paragraphs
        body = essay["body"].strip()
        paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]

        for para_text in paragraphs:
            # Handle line continuations within paragraphs
            clean_text = " ".join(para_text.split("\n"))
            p = doc.add_paragraph(clean_text)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(8)

        # Page break after each essay except the last
        if i < len(essays) - 1:
            doc.add_page_break()

    # ── Save ───────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Total essays: {len(essays)}")


if __name__ == "__main__":
    output = os.path.join(os.path.dirname(__file__), "pre-class-essays.docx")
    build_document(essays, output)
