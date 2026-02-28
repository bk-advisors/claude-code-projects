"""
Generate YouTube Upload Metadata for R for Management Consultants Videos
Output: youtube-metadata-intro-video.docx, youtube-metadata-full-lecture.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# -- Metadata Content -------------------------------------------------------

intro_video = {
    "title": "R for Management Consultants: Stop Copying, Start Coding | Course Intro",
    "description": (
        "A consultant spent three days rebuilding a quarterly report because two "
        "facilities were double-counted in her Excel file. Fourteen tabs, hundreds "
        "of linked formulas, and no way to trace what went wrong. That's the "
        "spreadsheet that fell apart — and the moment she learned R.\n\n"
        "This 5-minute introduction previews R for Management Consultants, a "
        "hands-on course that teaches R programming from scratch for global health "
        "professionals. No prior coding experience required — if you can use Excel, "
        "you can learn R.\n\n"
        "WHAT THIS COURSE COVERS:\n"
        "Nine chapters organized into four parts, taking you from your first line "
        "of R code to producing publication-quality reports for donors and program "
        "managers:\n\n"
        "Part I — Getting Started\n"
        "  • Chapter 1: Why R for Management Consultants\n"
        "  • Chapter 2: R Building Blocks (vectors, data types, functions)\n\n"
        "Part II — Working with Data\n"
        "  • Chapter 3: Importing & Exploring Data (CSV, Excel, DHIS2 exports)\n"
        "  • Chapter 4: Data Transformation with dplyr (filter, group, summarise)\n"
        "  • Chapter 5: Reshaping & Joining Datasets\n\n"
        "Part III — Visualization & Analysis\n"
        "  • Chapter 6: Data Visualization with ggplot2\n"
        "  • Chapter 7: Publication-Quality Charts\n"
        "  • Chapter 8: Descriptive Statistics for Consultants\n\n"
        "Part IV — Communication\n"
        "  • Chapter 9: Reproducible Reports with Quarto\n\n"
        "Every example uses real-world maternal and newborn health data — facility-level "
        "delivery data, neonatal mortality rates, equipment inventories, and training "
        "records.\n\n"
        "SERIES: Professional Skills for Global Health Consultants\n"
        "COMPANION COURSE: Inference and Intervention (causal models for MNH)\n"
        "PREREQUISITE: None — designed for complete beginners\n\n"
        "ABOUT THE SPEAKER:\n"
        "Matthew Kuch has spent over a decade in global child health — at CHAI in Uganda, "
        "Gavi in Geneva managing country portfolios across Kenya and Lesotho, and at Deloitte "
        "managing grants for USAID and DFID. He has watched brilliant analysts waste days "
        "on tasks that should take minutes and has seen teams transform their work by "
        "switching from spreadsheets to reproducible R workflows.\n\n"
        "BK Advisors | Professional Development for Global Health Consultants"
    ),
    "tags": (
        "R programming, management consulting, global health, data analysis, "
        "ggplot2, dplyr, tidyverse, maternal and newborn health, MNH, "
        "reproducible analysis, Quarto, data visualization, BK Advisors, "
        "Matthew Kuch, RStudio, Excel to R, data science for consultants"
    ),
    "category": "Education",
    "thumbnail_suggestions": (
        "TEXT OVERLAY: \"Excel → R\" in large white bold text with a stylized arrow, "
        "and subtitle \"Stop Copying. Start Coding.\"\n\n"
        "VISUAL: BK Advisors primary blue (#005CB9) background. A simplified code "
        "snippet or terminal icon visual element on the right side. "
        "\"BK ADVISORS\" branding in white at bottom-left.\n\n"
        "ALTERNATIVE: Split-screen concept — left side shows a messy Excel grid "
        "(faded, red-tinted), right side shows clean R code (bright, green-tinted). "
        "Arrow pointing from left to right. Text overlay: \"Stop Copying. Start Coding.\""
    ),
    "chapters": (
        "0:00 — The Spreadsheet That Fell Apart\n"
        "0:30 — About the Speaker\n"
        "1:00 — What This Course Covers\n"
        "2:00 — The Journey in Four Parts\n"
        "4:30 — Who This Is For\n"
        "5:00 — What You'll Walk Away With\n"
        "5:30 — Let's Get Started"
    ),
}

full_lecture = {
    "title": "R for Management Consultants | Complete Course — BK Advisors",
    "description": (
        "A complete, hands-on course that teaches R programming from scratch for "
        "management consultants, program analysts, and technical advisors in global "
        "health. Nine chapters, four parts, one goal: replace fragile spreadsheets "
        "with reproducible, auditable analyses that update with a single click.\n\n"
        "No prior programming experience is required. If you can use Excel, you can "
        "learn R. Every example uses real-world maternal and newborn health data.\n\n"
        "WHAT YOU'LL LEARN:\n\n"
        "PART I — GETTING STARTED\n"
        "Chapter 1: Why R for Management Consultants\n"
        "- When to use R vs. Excel, installing R + RStudio, your first commands\n"
        "- The case for reproducibility in consulting work\n\n"
        "Chapter 2: R Building Blocks\n"
        "- Vectors, data types, functions, installing and loading packages\n"
        "- The grammar of R: learning to read and write code\n\n"
        "PART II — WORKING WITH DATA\n"
        "Chapter 3: Importing & Exploring Data\n"
        "- read_csv, read_excel, tibbles, summary statistics, handling missing values\n"
        "- Working with DHIS2 exports and donor spreadsheets\n\n"
        "Chapter 4: Data Transformation with dplyr\n"
        "- The pipe operator, filter, select, mutate, summarise, group_by\n"
        "- Everything you do with pivot tables — but reproducible\n\n"
        "Chapter 5: Reshaping & Joining Datasets\n"
        "- pivot_longer, pivot_wider, left_join, binding rows\n"
        "- Combining data from multiple sources into one analysis\n\n"
        "PART III — VISUALIZATION & ANALYSIS\n"
        "Chapter 6: Data Visualization with ggplot2\n"
        "- The grammar of graphics: geoms, aesthetics, layers\n"
        "- Bar charts, line charts, scatter plots from your data\n\n"
        "Chapter 7: Publication-Quality Charts\n"
        "- Custom themes, faceting, multi-panel layouts with patchwork\n"
        "- Charts worthy of a donor presentation\n\n"
        "Chapter 8: Descriptive Statistics for Consultants\n"
        "- Means, proportions, confidence intervals, correlation, simple regression\n"
        "- The statistics every consulting report needs\n\n"
        "PART IV — COMMUNICATION\n"
        "Chapter 9: Reproducible Reports with Quarto\n"
        "- Code chunks, inline R, rendering to HTML, PDF, and DOCX\n"
        "- Parameterized reports: one template, multiple country briefs\n\n"
        "LEARNING OUTCOMES:\n"
        "By the end of this course, you will be able to:\n"
        "- Import and clean facility-level data from any CSV or Excel source\n"
        "- Transform raw data into summary tables and indicator dashboards\n"
        "- Create publication-quality charts with consistent branding\n"
        "- Calculate and interpret basic statistics with confidence intervals\n"
        "- Produce reproducible reports that update automatically when data changes\n"
        "- Build the R foundation for the Inference and Intervention course on causal models\n\n"
        "SERIES: Professional Skills for Global Health Consultants\n"
        "COMPANION COURSE: Inference and Intervention (causal models for MNH)\n"
        "PREREQUISITE LECTURES: Anatomy of a Good Point, Making Causal Arguments\n\n"
        "ABOUT THE SPEAKER:\n"
        "Matthew Kuch has spent over a decade in global child health — at CHAI in Uganda, "
        "Gavi in Geneva managing country portfolios across Kenya and Lesotho, and at Deloitte "
        "managing grants for USAID and DFID. He has watched brilliant analysts waste days "
        "on tasks that should take minutes and has seen teams transform their work by "
        "switching from spreadsheets to reproducible R workflows.\n\n"
        "BK Advisors | Professional Development for Global Health Consultants"
    ),
    "tags": (
        "R programming, management consulting, global health, data analysis, "
        "ggplot2, dplyr, tidyverse, maternal and newborn health, MNH, "
        "reproducible analysis, Quarto, data visualization, BK Advisors, "
        "Matthew Kuch, RStudio, Excel to R, data science for consultants, "
        "vectors, data frames, pivot_longer, pivot_wider, faceting, patchwork, "
        "confidence intervals, parameterized reports, read_csv, tibble, "
        "group_by, summarise, left_join, geom_bar, geom_line, geom_point, "
        "descriptive statistics, regression, reproducible reporting, "
        "DHIS2, facility data, program evaluation, donor reports"
    ),
    "category": "Education",
    "thumbnail_suggestions": (
        "TEXT OVERLAY: \"R for Consultants\" in large white bold text, "
        "with a \"COMPLETE COURSE\" badge in BK Advisors yellow (#FED141)\n\n"
        "VISUAL: Central \"R\" text surrounded by four coloured boxes "
        "representing the four parts of the course:\n"
        "  • \"Getting Started\" in light blue (#00A1DF)\n"
        "  • \"Data Wrangling\" in green (#83BD00)\n"
        "  • \"Visualization\" in orange (#FA7650)\n"
        "  • \"Communication\" in teal (#3E9B6E)\n"
        "Use BK Advisors primary blue (#005CB9) background. "
        "\"BK ADVISORS\" branding in white at bottom-left.\n\n"
        "ALTERNATIVE: Four-quadrant layout on blue background — each quadrant "
        "shows a simplified icon (terminal, table, chart, document) with the "
        "part name. Central \"R\" logo ties them together."
    ),
    "chapters": (
        "NOTE: Timestamps below are estimated at ~15 minutes per chapter.\n"
        "Adjust to match the actual video after reviewing.\n\n"
        "0:00:00 — Course Introduction\n"
        "0:03:00 — Chapter 1: Why R for Management Consultants\n"
        "0:18:00 — Chapter 2: R Building Blocks\n"
        "0:33:00 — Chapter 3: Importing & Exploring Data\n"
        "0:48:00 — Chapter 4: Data Transformation with dplyr\n"
        "1:03:00 — Chapter 5: Reshaping & Joining Datasets\n"
        "1:18:00 — Chapter 6: Data Visualization with ggplot2\n"
        "1:33:00 — Chapter 7: Publication-Quality Charts\n"
        "1:48:00 — Chapter 8: Descriptive Statistics for Consultants\n"
        "2:03:00 — Chapter 9: Reproducible Reports with Quarto\n"
        "2:18:00 — Course Wrap-Up & Next Steps"
    ),
}

# -- Document Builder -------------------------------------------------------


def build_document(metadata, video_filename, output_path):
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # -- Default paragraph font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # -- Heading styles --
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)  # BKA primary blue
        if level == 1:
            h.font.size = Pt(26)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(4)
        elif level == 2:
            h.font.size = Pt(16)
            h.font.italic = False
            h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        elif level == 3:
            h.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # -- Title Page --
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("YouTube Upload Metadata")
    run.font.name = "Calibri Light"
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
    run.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run(video_filename)
    run.font.name = "Calibri Light"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("BK Advisors")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)

    doc.add_page_break()

    # -- Video Title --
    doc.add_heading("Video Title", level=2)
    p = doc.add_paragraph(metadata["title"])
    p.paragraph_format.space_after = Pt(8)

    # -- Category --
    doc.add_heading("Category", level=2)
    p = doc.add_paragraph(metadata["category"])
    p.paragraph_format.space_after = Pt(8)

    # -- Description --
    doc.add_heading("Description", level=2)
    desc_paragraphs = metadata["description"].split("\n")
    for line in desc_paragraphs:
        if line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    # -- Tags --
    doc.add_heading("Tags", level=2)
    p = doc.add_paragraph(metadata["tags"])
    p.paragraph_format.space_after = Pt(8)

    # -- Video Chapters --
    doc.add_heading("Video Chapters", level=2)
    chapter_lines = metadata["chapters"].split("\n")
    for line in chapter_lines:
        if line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    # -- Thumbnail Suggestions --
    doc.add_heading("Thumbnail Suggestions", level=2)
    thumb_paragraphs = metadata["thumbnail_suggestions"].split("\n")
    for line in thumb_paragraphs:
        if line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    # -- Save --
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    base_dir = os.path.dirname(__file__)

    build_document(
        intro_video,
        "Intro - R for Management Consultants.mp4",
        os.path.join(base_dir, "youtube-metadata-intro-video.docx"),
    )

    build_document(
        full_lecture,
        "Complete Course - R for Management Consultants.mp4",
        os.path.join(base_dir, "youtube-metadata-full-lecture.docx"),
    )
