"""
Generate Pre-Class Essays for "Inference and Intervention" Course
Output: pre-class-essays.docx
Style: Morgan Housel (conversational, story-driven, counterintuitive insights)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

# ── Essay Content ──────────────────────────────────────────────────────────────

essays = []

# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 0,
    "title": "The Gap Between Good Intentions and Good Outcomes",
    "subtitle": "A roadmap for the intellectual journey ahead",
    "body": r"""
In 2015, a group of philanthropists pooled hundreds of millions of dollars to save newborn lives across sub-Saharan Africa. They hired the best technical advisors. They partnered with governments. They funded training, equipment, community outreach, and data systems. Their intentions were as good as intentions get.

Five years later, some countries showed dramatic improvements. Others showed almost none. And in at least one case, the investment appeared to have made things worse — not because the programs were harmful, but because the government had quietly reduced its own health spending, effectively using donor money to subsidize other priorities. The net impact was a fraction of what the models predicted.

This is not an unusual story. It is, in many ways, *the* story of global health investment over the past three decades. Enormous resources. Genuine commitment. Mixed results. And a persistent, uncomfortable question: why is the gap between what we spend and what we achieve so stubbornly large?

The answer, I believe, is not about money, or politics, or even corruption — though all of those play a role. The answer is about how we think. Specifically, how we think about cause and effect.

Most decision-makers in global health are trained to analyze data. They can run regressions, build dashboards, and calculate cost-effectiveness ratios. What they're rarely trained to do is think causally — to ask not just "what happened?" but "why did it happen?" and "what would happen if we did something different?"

These sound like the same question. They are not. And the distance between them is where this course lives.

Let me give you a preview of the journey.

We begin with a deceptively simple idea: before you touch any data, draw a picture of how you think the world works. Not a flowchart. Not an org chart. A causal model — a diagram with nodes and arrows that represents your theory about what causes what.

This sounds almost childishly simple. Draw some circles and arrows on a whiteboard. But the act of drawing forces a discipline that most analytical work skips entirely. It forces you to state your assumptions explicitly. It forces you to decide: does training cause quality, or does quality attract trained staff? Does government funding drive equipment purchases, or do equipment needs drive government funding? The arrow has to point one way or the other. You can't hide behind ambiguity.

In the first three chapters — what I call the Foundations — you'll learn to think this way. Chapter 1 introduces the core distinction between observing a situation and intervening in it, and shows you the three traps that catch analysts who skip the causal model: confounding, reverse causation, and selection bias. Chapter 2 teaches you the grammar of causal diagrams — the three fundamental structures that govern how information flows. Chapter 3 puts you in the field, building a real causal model through iterative stakeholder interviews, watching it grow from four nodes to twelve as each conversation reveals what the last one missed.

By the end of Part I, you'll have a qualitative causal model — a picture. It captures the structure of your beliefs, but not their strength. You believe training affects quality, but how much? You believe government commitment drives funding, but with what probability?

Part II — Chapters 4 through 6 — adds the numbers. Chapter 4 introduces probability, conditional probability tables, Bayes' Rule, and the Causal Markov Condition that ties it all together into a Bayesian network. Chapter 5 shows you how to use that network for diagnosis — running the model backwards from observed outcomes to infer probable causes, including the counterintuitive phenomenon of "explaining away," where confirming one cause reduces your belief in alternatives. Chapter 6 drops you into the real world of cost-effectiveness analysis, where you'll encounter Simpson's Paradox — the unsettling fact that aggregated data can recommend the exact opposite of what disaggregated data shows — and learn why only the causal model can tell you which answer is right.

Part II is where most analytics courses end. You have a model, you have numbers, you can compute expectations. But here's the thing: computing an expected value is not the same as making a decision. And making a decision is not the same as making a decision in a world where other people are also making decisions.

Part III — Chapters 7 through 9 — is where the course gets real.

Chapter 7 introduces the do-operator: the formal distinction between observing that something is true and making it true. This is where you learn graph surgery — the technique of cutting arrows in your causal diagram to compute the effect of an intervention rather than the effect of an observation. It's the difference between noticing that hospitals with more equipment have better outcomes and actually buying equipment for a hospital that doesn't have it. The former is contaminated by confounders. The latter is a clean causal effect.

Chapter 8 takes you from single decisions to portfolio allocation. You have a budget. You have ten countries. Each country has different costs, different uncertainties, different government dynamics. How do you allocate? The same logic that Harry Markowitz used to revolutionize financial investing — diversify, quantify risk, think in terms of portfolios rather than individual bets — applies directly. You'll run Monte Carlo simulations, compute risk-return tradeoffs, and learn to think in terms of sequential decisions: commit a little, observe results, then scale.

Chapter 9 adds the hardest layer: other people have strategies too. When you invest in a country, the government responds — sometimes by co-investing (crowding in), sometimes by redirecting its own funds elsewhere (crowding out). This is game theory, and it transforms the resource allocation problem from an optimization puzzle into a strategic interaction. You'll learn about Nash Equilibrium, the free-rider problem, and — most importantly — how to design contracts and commitment devices that make cooperation the rational choice.

Finally, Chapter 10 closes the loop. You've been building models from expertise and interviews. But can data build models too? The answer is a qualified yes — algorithms can discover some causal structures from observational data, but they cannot distinguish between all structures. Some arrows remain fundamentally ambiguous without expert knowledge. The chapter teaches you how to combine algorithmic discovery with domain expertise, and why neither alone is sufficient.

That's the journey. From a correlation to a causal model. From a picture to a number. From an observation to an intervention. From a single decision to a portfolio. From a solitary optimizer to a strategic actor. And from intuition to a formal, updatable, testable framework for thinking about cause and effect.

Why does this matter?

Because every year, billions of dollars are allocated to global health programs based on analyses that confuse correlation with causation, that ignore strategic responses by governments, that aggregate data in ways that reverse the truth, and that treat observation as if it were intervention. These are not rare mistakes made by careless analysts. They are systematic errors embedded in the standard analytical toolkit — errors that only a causal framework can correct.

The economist John Maynard Keynes reportedly said, "It is better to be roughly right than precisely wrong." The tools in this course won't make you precisely right — the world is too uncertain for that. But they will keep you from being precisely wrong. They will give you a structured way to think about complex systems, to communicate your assumptions transparently, to update your beliefs as evidence accumulates, and to make decisions that account for uncertainty, strategic behavior, and the fundamental difference between watching and doing.

That's a lot to promise from a course. But here's the thing about causal thinking: once you learn to see it, you can't unsee it. Every regression table, every program evaluation, every resource allocation debate — they all look different once you have a causal model in your head. Not because the model gives you the answer. But because it shows you what the answer depends on.

And knowing what the answer depends on is the beginning of wisdom.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 1,
    "title": "The Most Dangerous Number Is a Correlation",
    "subtitle": "Why trusting data without a causal model can lead you to recommend the worst possible option",
    "body": r"""
In 2019, a global health organization faced a decision that would affect hundreds of thousands of newborns across sub-Saharan Africa. They had two competing programs — let's call them Feather Touch and TruSmartz — and a limited budget. So they did what any rigorous organization would do. They hired consultants. They ran regressions. They built dashboards. And the data spoke clearly: Feather Touch was the winner.

There was just one problem. The data was lying.

Not intentionally, of course. Data doesn't have intentions. But the correlation between Feather Touch's presence and better outcomes was driven entirely by a third factor the analysts never measured. The places where Feather Touch operated happened to be the places that were already improving for other reasons — better infrastructure, more government attention, stronger local health systems. Feather Touch wasn't causing success. It was surfing on it.

TruSmartz, meanwhile, had been deployed in the hardest districts. The places no one else wanted to go. The facilities with the worst baseline conditions. When you ran a simple regression, TruSmartz looked terrible. Its numbers were dragged down by the very difficulty of the environments it served. The analysts saw the negative coefficient and moved on.

This is the story of how a well-funded organization nearly doubled down on the wrong program — and it happens far more often than anyone in the development sector wants to admit.

Here's what makes this story unsettling: the analysts weren't incompetent. They were doing exactly what they'd been trained to do. They collected data, they ran statistical models, they reported findings. The problem wasn't the math. The problem was the absence of a causal model — a story about *why* things happen, not just *what* happens alongside *what*.

There's a famous example in medicine. For decades, doctors noticed that patients who took hormone replacement therapy (HRT) had lower rates of heart disease. The correlation was strong and consistent. Millions of women were prescribed HRT partly on this basis. Then randomized trials showed that HRT actually *increased* heart disease risk. The correlation had been driven by the fact that women who chose HRT tended to be wealthier, more health-conscious, and had better access to medical care. The treatment was taking credit for the patient's background.

The same pattern plays out everywhere, but it's especially dangerous in global health, where the stakes are measured in lives and the data is messy.

When we talk about "causal thinking," we're really talking about a discipline — a habit of mind that forces you to articulate *how* you think the world works before you start interrogating data. It's the difference between asking "what correlates with what?" and asking "what causes what, and through which pathways?"

Consider three traps that catch even sophisticated analysts.

The first is confounding. This is the Feather Touch problem. Two things move together not because one causes the other, but because a third factor drives both. A country that invests heavily in health worker training might simultaneously invest in facility equipment. If you measure the effect of equipment without accounting for training, you'll overestimate what equipment alone can do. The confounder — government commitment — inflates the apparent effect.

The second is reverse causation. Does having more doctors in a district cause better health outcomes, or do districts with better outcomes attract more doctors? Does foreign aid improve governance, or do better-governed countries attract more aid? The arrow of causation matters enormously, but in cross-sectional data, it's invisible.

The third is selection bias. This is perhaps the sneakiest. Imagine you're studying which characteristics predict success among health facilities that received a grant. You survey the grant recipients and find that facilities with more experienced staff perform better. Seems obvious. But you're only looking at facilities that *got* the grant. The selection process itself may have filtered for certain combinations of characteristics, creating correlations that don't exist in the broader population. You've conditioned on an outcome, and now your world looks different from reality.

These three traps share something in common: they all produce real, reproducible statistical patterns that point in the wrong direction. The correlations are not artifacts. They're not noise. They're genuine patterns in the data — they just don't mean what you think they mean.

This is why the Feather Touch story matters. It's not a story about bad data or lazy analysis. It's a story about what happens when you skip the hardest step in analytical reasoning: building a causal model before you touch the data.

A causal model is, at its core, a drawing. It's a diagram that says: "I believe A affects B, B affects C, and D affects both A and C." It forces you to be explicit about your assumptions. And once those assumptions are on paper, you can interrogate them. You can ask: "If this model is right, what patterns should I see in the data? What patterns should I *not* see?" You can design your analysis to distinguish between your model and alternatives.

Without that discipline, you're just reporting correlations and hoping they mean something. And as the Feather Touch story shows, hope is not a strategy when lives are at stake.

The psychologist Daniel Kahneman once observed that humans are "a machine for jumping to conclusions." We see patterns and immediately construct narratives to explain them. We're wired to confuse sequence with consequence, correlation with causation, association with explanation. It takes deliberate effort to resist this instinct — to slow down and ask, "What else could explain this pattern?"

That's what this course is about. Not statistics. Not software. Not even causal diagrams, though you'll learn to build those. It's about developing a disciplined way of thinking about cause and effect in complex systems — the kind of systems where millions of dollars and millions of lives hang in the balance.

The most dangerous number isn't a wrong number. It's a right number with the wrong interpretation. And the only defense against misinterpretation is a causal model that tells you what to look for — and what to look out for.

In the chapters ahead, we'll build that model piece by piece. We'll start with qualitative sketches — just arrows and nodes — and gradually add probabilities, decisions, and strategic interactions. By the end, you'll have a framework for thinking about maternal and newborn health investments that goes far beyond regression tables and dashboards.

But it all starts here, with a simple commitment: before you trust any number, draw the picture of how you think the world works. Then check whether the number still tells the same story.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 2,
    "title": "Drawing What You Believe",
    "subtitle": "Why writing down your causal beliefs is the most productive argument you'll ever have",
    "body": r"""
In any room where important decisions are being made about health programs, there are at least three causal models operating simultaneously. None of them are written down.

The epidemiologist believes that training health workers is the lever that moves outcomes. The economist believes it's about incentives and resource allocation. The program manager believes it's about supply chains and equipment availability. Each person has a mental model of how the world works — a set of beliefs about what causes what — and each person's model leads to a different recommendation.

The tragedy isn't that they disagree. The tragedy is that they don't even know they disagree. Because their models live inside their heads, the conversation happens at the level of conclusions rather than assumptions. They argue about what to do without ever surfacing *why* they think different things will work. It's like three people arguing about directions while looking at three different maps.

This is the problem that qualitative causal models solve. Not by telling you who's right, but by forcing everyone to draw their map — and then comparing them.

A causal model, in its simplest form, is a diagram. It has nodes — things you care about — and arrows — your beliefs about what influences what. That's it. No equations. No data. Just a visual representation of your theory about how the world works.

But simplicity is deceptive. The act of drawing this diagram turns out to be one of the most powerful analytical exercises you can do. Here's why.

When your beliefs live inside your head, they enjoy a kind of diplomatic immunity. They don't have to be consistent. They don't have to be complete. You can simultaneously believe that "training improves quality" and "quality doesn't affect outcomes" without ever noticing the contradiction, because you never think about both beliefs at the same time.

The moment you draw an arrow from "Training" to "Quality of Care" and then try to connect "Quality of Care" to "Neonatal Mortality," the contradiction becomes visible. Either the arrow is there or it isn't. Either you believe training matters for outcomes — through quality — or you don't. The diagram doesn't let you hide.

There are four types of nodes in a causal model, and they matter more than you might think.

The first is a probabilistic node — something uncertain that you can't directly control. "Will the government co-finance this program?" You don't know. It depends on politics, budget cycles, competing priorities. You can influence it, maybe, but you can't set it. These are the oval-shaped nodes, and they represent the uncertainty that makes decision-making hard.

The second is an objective node — the thing you're ultimately trying to affect. "Neonatal deaths averted." "Cost per life saved." This is your hexagon, your destination. Everything else in the model exists in relation to this node.

The third is a strategic option — something you *can* control. "Invest in CPAP machines" or "Fund community health worker training." These are your rectangles, your levers. The whole point of building the model is to figure out which levers to pull.

The fourth — and most often overlooked — is a function node. This is a quantity that is completely determined by its inputs. If you know the number of facilities and the cost per facility, you know the total cost. There's no uncertainty here, just arithmetic. These chevron-shaped nodes seem boring, but they keep the model honest by separating what's uncertain from what's just calculation.

Now, the arrows. An arrow from A to B means you believe A has a causal influence on B. Not a correlation. Not an association. A genuine causal effect: if you could reach in and change A, B would change as a result. The sign on the arrow — positive or negative — tells you the direction. More training leads to better quality (positive). Higher cost leads to fewer facilities funded (negative).

Here's where it gets interesting. The way nodes connect to each other creates structures that have profound implications for what you can learn from data. There are exactly three fundamental structures, and understanding them is like learning the grammar of causation.

The first is the serial chain: A causes B, and B causes C. Think of a domino effect. Government funding leads to equipment purchases, which lead to better care quality. In this structure, A and C are correlated — you'll see them move together in data — but only because B carries the influence. If you could somehow hold B constant, A and C would become independent. The middleman is doing all the work.

The second is the diverging fork: B causes both A and C. Think of a common cause. A country's overall health system strength affects both its training quality and its equipment availability. In data, training and equipment will be correlated — not because one causes the other, but because they share a common parent. Again, control for B, and the correlation vanishes.

The third structure is the one that breaks most people's intuition: the converging collider. A and B both cause C, but A and B are otherwise unrelated. Think of two independent factors — say, staff motivation and equipment quality — that both contribute to patient outcomes. In the general population, motivation and equipment are unrelated. But here's the twist: if you only look at facilities with good outcomes (you condition on C), suddenly motivation and equipment become negatively correlated. Among successful facilities, the ones that succeeded despite poor equipment must have had exceptional staff, and vice versa. By selecting on the outcome, you've created a phantom correlation between its causes.

This collider effect is responsible for an enormous number of analytical errors. It's why studying only successful companies gives you misleading predictors of success. It's why looking only at admitted hospital patients can make diseases appear related when they're not. It's why conditioning on the wrong variable doesn't just fail to help — it actively misleads.

These three structures — chains, forks, and colliders — are the atoms of causal reasoning. Every complex causal model, no matter how many nodes and arrows it contains, is built from combinations of these three patterns. If you understand how information flows (and gets blocked) in each one, you understand the logic of causation.

And this brings us back to the room full of experts who can't agree. When you ask each person to draw their causal model — to literally put nodes and arrows on a whiteboard — something remarkable happens. The disagreements become productive. Instead of arguing about conclusions, people argue about arrows. "Do you really think government commitment directly affects neonatal mortality, or does it work through facility funding?" "Is equipment quality a cause of staff retention, or is it the other way around?"

These are arguments you can resolve with evidence, logic, or at least structured discussion. They're a thousand times more useful than arguing about whether to fund Program A or Program B.

The philosopher of science Karl Popper said that the point of articulating a theory is to make it falsifiable — to put it in a form where it can be tested and potentially proven wrong. A causal diagram does exactly that. It transforms vague intuitions into specific, testable claims about the structure of the world.

You don't need data to draw a causal model. You need honesty. You need the willingness to say, "Here's what I believe, and here's where I'm uncertain." The data comes later — to test, refine, and quantify the model. But the model comes first.

Draw what you believe. Then be prepared to be wrong. That's how understanding begins.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 3,
    "title": "The Art of Asking Better Questions",
    "subtitle": "How each conversation reveals what the previous model couldn't see",
    "body": r"""
There's a moment in every consulting engagement — whether you're advising a Fortune 500 company or a Ministry of Health — when you realize that your first model was embarrassingly incomplete. Not wrong, exactly. Just missing the things that actually matter.

This moment is not a failure. It's the point.

Consider a team trying to understand why neonatal mortality varies so dramatically across districts in a sub-Saharan African country. They start where any reasonable team would start: with the obvious variables. Facility readiness. Equipment availability. Staff training. These are the things that appear in WHO reports, the things that donors measure, the things that everyone agrees matter.

So they draw a model with four nodes. Equipment quality affects care quality. Staff training affects care quality. Care quality affects neonatal outcomes. Facility readiness sits underneath it all. It's clean. It's logical. It fits on a single whiteboard.

Then they walk into the Ministry of Health and sit down with the Director of Maternal and Child Health.

Twenty minutes into the conversation, the model is already in trouble. The Director mentions something the team hadn't considered: district management capacity. "You can send all the equipment you want," she says, "but if the District Health Management Team can't maintain it, it'll be sitting in a corner gathering dust within six months."

She mentions another thing: community health worker networks. "Mothers don't just show up at facilities. Someone has to convince them to come. Someone has to follow up after delivery. If your community health workers are overwhelmed or unmotivated, your facility could be perfect and you'd still see poor outcomes."

The team goes back to their whiteboard. The four-node model becomes an eight-node model. They add district management, community outreach, equipment maintenance, and referral systems. The new model has arrows the team hadn't imagined — district management affecting not just equipment but also staff morale, community outreach affecting both facility utilization and early detection of complications.

Better. But still not right.

For the second round, they visit three facility directors in different districts — one high-performing, one average, one struggling. The conversations are revelatory in different ways.

The high-performing facility director talks about something unexpected: staff retention. "My nurses are good because they've been here for years. They know the community. They know the referral pathways. When I lose a nurse to the capital city, it takes two years to rebuild that knowledge." The team adds a node for staff retention and draws arrows from salary levels, housing conditions, and workload to retention, and from retention to care quality.

The struggling facility director tells a different story. She talks about stockouts — not just of medication, but of basic supplies like gloves and syringes. "We can't provide quality care when we're rationing supplies," she says. And then she adds something that restructures the entire model: "The stockouts aren't random. They happen because our district health team doesn't submit orders on time. And they don't submit orders on time because they're understaffed and overwhelmed with reporting requirements."

Suddenly, what looked like a supply chain problem is actually a governance problem. The arrow doesn't go from "national supply chain" to "stockouts." It goes from "district management capacity" to "order timeliness" to "supply availability" to "care quality." The causal pathway is longer and more institutional than the team initially imagined.

The average facility director provides perhaps the most surprising insight: "You're asking about equipment and training, but the biggest factor in my outcomes is whether the mother arrives in time. Half my complicated cases arrive too late because they tried a traditional birth attendant first, or because there's no transport at night. The facility side is only half the equation."

Round three takes the team to the community level. They interview community health workers — the people closest to the mothers and families. And here the model expands again in ways no one sitting in the capital could have predicted.

A community health worker in a rural district explains: "I'm supposed to visit every pregnant woman in my catchment area four times during pregnancy. But I have 800 households. I don't have a motorcycle. Some of these homes are three hours' walk. I visit who I can, and the rest I pray for."

Another describes the social dynamics: "Even when I reach a mother and tell her to deliver at the facility, her mother-in-law often overrides me. The mother-in-law delivered all her children at home, and she doesn't trust the facility. I'm fighting culture, not just distance."

The model now has twelve or fifteen nodes. It includes things like transport availability, cultural attitudes toward facility-based delivery, community health worker workload, geographic accessibility, and mother-in-law influence. It's messy. It's complicated. And it's much, much closer to reality than the neat four-node model the team started with.

This iterative process — build a model, test it against reality, revise, repeat — is the heart of causal analysis. And it illustrates something that Morgan Housel might call the paradox of expertise: the more you learn about a system, the less simple it looks, but the better your decisions become.

There's a temptation, especially among quantitative analysts, to skip the interview stage and go straight to the data. After all, we have DHIS2 databases with millions of records. We have DHS surveys. We have facility assessments. Why spend weeks talking to people when we could spend days running regressions?

The answer is that data can tell you *what* is correlated with *what*, but it cannot tell you *why*. It cannot tell you that the reason equipment breaks down is that district managers don't submit maintenance requests. It cannot tell you that mothers arrive late because of transport, not ignorance. It cannot tell you that community health workers are spread impossibly thin.

More importantly, data cannot tell you what's missing from your model. You can't include a variable in your regression if you don't know it exists. And you can't know it exists if you haven't talked to the people who live inside the system.

Every interview round in this process does two things. First, it adds new nodes and arrows — new variables and new causal pathways that the team hadn't considered. Second, and more subtly, it changes the *structure* of existing relationships. An arrow that the team drew as direct (equipment → outcomes) turns out to be mediated by maintenance capacity. A relationship they assumed was universal turns out to vary by context.

This is the art of asking better questions. Not "What are the determinants of neonatal mortality?" — that's too broad. But "You mentioned that equipment breaks down. Can you walk me through what happens when a piece of equipment malfunctions? Who do you call? How long does it take to get fixed?" These granular, process-oriented questions reveal the causal mechanisms that aggregate data hides.

The novelist John Steinbeck once wrote, "The discipline of the written word punishes both stupidity and dishonesty." The same is true of causal diagrams. Each interview punishes the model's incompleteness, forcing it to grow, to adapt, to accommodate the messy reality of how health systems actually function.

Your first model will be wrong. Your second model will be less wrong. Your third model will be less wrong still. The goal isn't perfection — it's progressive honesty about the system you're trying to change.

And the key to that honesty? Talking to the people who know the system better than you ever will, and having the humility to redraw your model when they tell you something you didn't expect.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 4,
    "title": "Putting Numbers on Uncertainty",
    "subtitle": 'Why "I don\'t know" is the beginning of analysis, not the end',
    "body": r"""
There's an old joke among statisticians: "All models are wrong, but some are useful." It's attributed to George Box, and it's usually deployed as a kind of intellectual shrug — an acknowledgment that perfection is impossible, so let's just do our best.

But there's a deeper truth hiding inside that quip. The useful models aren't the ones that are less wrong. They're the ones that are *honestly* wrong — the ones that explicitly state their uncertainty rather than hiding it behind a veneer of false precision.

In the previous chapters, we built qualitative causal models — diagrams with arrows showing what causes what. Those models captured the structure of our beliefs. Now comes the harder part: putting numbers on those beliefs. Not because the numbers will be right, but because the act of quantifying forces a precision that qualitative reasoning lets you avoid.

Here's what I mean. You can draw an arrow from "Staff Training" to "Quality of Care" and nod along. Yes, training matters. Everyone agrees. But now try to answer this: if a facility has well-trained staff, what's the probability that care quality is high? Is it 90%? 70%? 50%? And if staff are poorly trained, what's the probability of high quality? 10%? 30%?

These questions feel uncomfortable. They should. They're asking you to commit to a specific belief about the strength of a causal relationship. And the moment you commit, your belief becomes testable. It can be compared to data, challenged by colleagues, and updated when new evidence arrives.

This is the power of probability — not as a mathematical abstraction, but as a language for expressing uncertainty.

Let's start with the basics. A probability is a number between 0 and 1 that represents your degree of belief that something is true. P(A) = 0.7 means you believe there's a 70% chance that A is the case. It's not a frequency (though it can be informed by frequencies). It's a statement about your state of knowledge.

Conditional probability — P(A|B) — is where things get interesting. This is your belief about A, given that you know B is true. P(High Quality | Well Trained) = 0.8 means that if you know staff are well-trained, you believe there's an 80% chance that care quality is high. This is different from P(High Quality) on its own, which averages over all possible training levels.

A Conditional Probability Table, or CPT, takes this idea and systematizes it. For each node in your causal model, you specify the probability of each state of that node for every combination of its parent states. If "Quality of Care" has two parents — "Staff Training" (high/low) and "Equipment Condition" (good/poor) — then you need to fill in four probabilities: P(High Quality | High Training, Good Equipment), P(High Quality | High Training, Poor Equipment), P(High Quality | Low Training, Good Equipment), and P(High Quality | Low Training, Poor Equipment).

This feels tedious. It is. But consider what you get in return. You now have a complete, internally consistent specification of how you believe training and equipment interact to produce quality. Maybe you believe that good equipment partially compensates for poor training — the probability of high quality is 0.5 even with low training if equipment is good. Or maybe you believe that both must be present — the probability drops to 0.1 if either factor is deficient. These are genuinely different beliefs with genuinely different policy implications, and the CPT forces you to choose.

Now, Bayes' Rule. If there's one formula that changes how you see the world, this is it.

Here's the intuition. You're a district health officer. You receive a report that a facility's CPAP machine has malfunctioned. You want to know: is the problem due to poor maintenance, or is the machine itself defective? You have some prior beliefs — maybe 20% of CPAPs in your district have manufacturing defects, while 40% of facilities have inadequate maintenance protocols. You also know something about how these causes produce the symptom: defective machines malfunction 80% of the time, while poorly maintained machines malfunction 50% of the time.

Bayes' Rule tells you how to combine these pieces. It says: the probability that the cause is defective manufacturing, given that you've observed a malfunction, is proportional to the prior probability of a defect times the probability of malfunction given a defect. You compute the same thing for each possible cause, normalize so the probabilities add to 1, and you have your updated beliefs.

What makes Bayes' Rule profound is not the math — it's the logic. It tells you that diagnosis is not about finding the most likely cause in isolation. It's about weighing each cause by both its prior plausibility *and* its ability to explain what you've observed. A rare cause that perfectly explains the evidence can beat a common cause that weakly explains it.

This brings us to the Causal Markov Condition — the bridge between causal models and probability. It says: given its direct causes (its parents in the causal diagram), a variable is independent of everything else that isn't its descendant. In plain English: once you know a node's immediate causes, knowing about its grandparents or distant cousins tells you nothing additional.

This seems like a technical detail, but it's actually a gift. It means you can build a complex model with dozens of nodes and still compute probabilities efficiently. Instead of needing one enormous probability table for all variables simultaneously, you only need small local tables — one per node — that capture each variable's relationship with its direct parents. The Causal Markov Condition guarantees that these local tables, combined with the structure of the graph, fully determine the global probability distribution.

This is what a Bayesian network is: a causal diagram plus a set of CPTs, one per node. The diagram tells you the qualitative story — what causes what. The CPTs tell you the quantitative story — how strongly, and with what probability.

Building a Bayesian network is an exercise in structured humility. You don't need to know everything. You just need to state, for each variable, what you believe about its relationship with its direct causes. "I don't know" is not an acceptable answer — but "I think there's roughly a 60% chance, give or take" absolutely is. The model doesn't require certainty. It requires honesty about your uncertainty.

And here's the payoff: once you've built the network, it can do things your intuition cannot. It can propagate evidence through chains of variables, updating beliefs across the entire model when you observe a single fact. It can handle explaining away — the phenomenon where observing one cause makes alternative causes less likely. It can identify which pieces of information would be most valuable to collect. It can even compute the expected value of gathering additional data before making a decision.

None of this is possible with a qualitative diagram alone. The arrows tell you the shape of your beliefs, but the numbers make those beliefs actionable.

In the chapters ahead, we'll use these Bayesian networks to diagnose district-level health systems, evaluate intervention strategies, and allocate resources across countries. The numbers won't be perfect — they never are. But they'll be explicit, updatable, and far more useful than the vague intuitions they replace.

The philosopher Bertrand Russell once said, "The trouble with the world is that the stupid are cocksure and the intelligent are full of doubt." In causal analysis, doubt isn't a weakness. It's a feature. The trick is to measure it.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 5,
    "title": "Running the Model Backwards",
    "subtitle": "How observing an outcome changes everything you think about its causes",
    "body": r"""
Something strange happens when you learn that a patient has recovered from a severe illness. Your beliefs about what caused the illness — and what treatment was responsible for the recovery — should change. And they do change, but not always in the direction you'd expect.

This is the phenomenon of "explaining away," and it is one of the most counterintuitive and important ideas in probabilistic reasoning. Once you see it, you can't unsee it — and you'll start noticing it everywhere.

Let's start with a simple example. Imagine two possible reasons a newborn might survive a complicated delivery: the facility had a functioning CPAP machine, or the attending midwife had advanced resuscitation training. Before you know the outcome, these two factors are independent — knowing whether the facility has CPAP tells you nothing about whether the midwife has training. They're separate investments, separate programs, separate budget lines.

Now suppose you learn that the baby survived. What happens to your beliefs?

If you then learn that the facility had a functioning CPAP, your belief that the midwife had advanced training goes *down*. Not because CPAP somehow prevents training — that would be absurd. But because the survival has been "explained" by the CPAP. The need to invoke a second explanation is reduced. The survival is no longer a mystery that requires multiple contributing factors.

This is explaining away. It creates a negative correlation between causes *after you condition on their common effect*. Before you know the outcome, the causes are independent. After you know the outcome, they compete to explain it.

If this feels weird, you're in good company. Explaining away violates a deep intuition that learning about one cause shouldn't affect your beliefs about an unrelated cause. But the key word is "unrelated." Once they share a common effect, and you observe that effect, they're no longer unrelated. They're connected through the outcome, like two suspects in a crime who become linked the moment you know a crime was committed.

This idea is not just an intellectual curiosity. It has profound practical implications for how we diagnose problems and evaluate programs in global health.

Consider a district health officer trying to understand why neonatal mortality has spiked in her district. She has a causal model with several potential explanations: supply chain disruptions, staff turnover, decreased community outreach, and seasonal disease patterns. Before investigating, she treats these as somewhat independent hypotheses.

She starts gathering evidence. First, she confirms that there was indeed a major supply chain disruption — several facilities ran out of essential medications for weeks. This explains a lot. Her belief in supply chain disruption as the primary cause goes up sharply.

But notice what else happens. Her belief in staff turnover as a contributing cause goes *down*. Not because she has evidence against it — she hasn't checked yet. But because the supply chain disruption already accounts for much of the observed increase. The spike has been partially "explained away."

This is how Bayesian updating works in a converging structure — a collider where multiple causes feed into a common effect. Observing the effect activates the connection between causes. Confirming one cause diminishes the others. This is not a bias or an error. It's the correct probabilistic reasoning given the model structure.

Now let's apply this to a full situational analysis.

Imagine you're assessing two districts: Mwanga and Songea. Both have concerning neonatal mortality rates, but you suspect the underlying causes are different. Your Bayesian network includes nodes for government commitment, facility funding, equipment quality, staff training, care quality, and neonatal outcomes.

For District Mwanga, you observe that equipment quality is poor but staff training is adequate. You enter this evidence into your model and hit "update." The model propagates your evidence in both directions — forward to outcomes and backward to causes.

Forward propagation is intuitive: poor equipment leads to lower expected care quality, which leads to worse expected outcomes. No surprises there.

But backward propagation is where the insight lives. Given that equipment is poor, what do you now believe about government commitment? If your model says government commitment causes both equipment quality and facility funding, then observing poor equipment shifts your beliefs about government commitment downward. And since government commitment also affects staff training, you might expect to see poor training too.

But wait — you observed that training is *adequate*. This is surprising given your updated belief about government commitment. The model now has to reconcile two pieces of conflicting evidence about government commitment: equipment quality (suggesting low commitment) and staff training (suggesting higher commitment). The result is a nuanced posterior belief — government commitment is probably moderate, or perhaps commitment is present but poorly directed.

This kind of reasoning — integrating multiple pieces of evidence, some confirming and some contradicting, through a causal structure — is something humans do poorly in their heads. We tend to anchor on the most vivid or recent evidence and ignore how it interacts with other observations. The Bayesian network does it correctly and consistently.

For District Songea, the evidence pattern is different. Equipment is adequate, but community outreach has collapsed. The model traces this back to different root causes — perhaps a reorganization of the community health worker program, or a funding shift that prioritized facility-based care at the expense of community engagement. The diagnosis is fundamentally different, and therefore the recommended intervention is different too.

This is the power of situational analysis: the same model, applied to different evidence, produces different diagnoses and different recommendations. One size does not fit all. The model adapts to the local context because it processes evidence through a causal structure rather than applying a universal formula.

There's a broader lesson here about how we process evidence in everyday life. We're all running models backwards all the time, usually without realizing it. When a colleague gets promoted, we update our beliefs about their competence — and simultaneously (through explaining away) reduce our suspicion that it was just politics. When a country achieves rapid economic growth, we give credit to its latest policy reform — and explaining away reduces the credit we give to favorable global conditions.

The problem is that we do this intuitively and inconsistently. We explain away too aggressively when we like one explanation, and not aggressively enough when we don't. We update on vivid evidence and ignore subtle evidence. We forget that observing an outcome changes the relationship between its causes.

A formal model doesn't have these biases. It updates beliefs exactly as the laws of probability dictate. It doesn't care which explanation is politically convenient or emotionally satisfying. It just follows the arrows and the numbers.

This doesn't mean the model is always right. It's only as good as its structure and its probabilities. But it provides a disciplined, transparent baseline that human judgment can then refine and contextualize. It's a starting point for reasoning, not a substitute for it.

The detective novelist Raymond Chandler once wrote, "When in doubt, have a man come through a door with a gun in his hand." It's great advice for fiction. In causal analysis, the equivalent is: when in doubt, enter the evidence you have and let the model tell you what it implies — not just about the evidence itself, but about everything connected to it.

Run the model backwards. You might be surprised by what you learn about the causes you haven't observed yet.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 6,
    "title": "When the Average Lies",
    "subtitle": "Why aggregated data can recommend the exact wrong thing — and how to catch it",
    "body": r"""
In 1973, the University of California, Berkeley was sued for gender discrimination in graduate admissions. The numbers looked damning: 44% of male applicants were admitted, compared to only 35% of female applicants. The gap was statistically significant. The case seemed open and shut.

Then a statistician named Peter Bickel did something that changed the narrative entirely. He looked at the data department by department. And in most individual departments, women were admitted at *higher* rates than men. The overall gender gap wasn't caused by discrimination within departments. It was caused by women disproportionately applying to more competitive departments with lower admission rates for everyone.

This is Simpson's Paradox — the phenomenon where a trend that appears in aggregated data reverses when the data is broken into meaningful subgroups. It's not a statistical anomaly or a rare edge case. It happens all the time, in every field, and it has led to some of the worst decisions in the history of medicine, law, and public policy.

The reason Simpson's Paradox matters for our purposes is that global health data is riddled with it.

Consider this scenario. You're evaluating a CPAP program across two countries. In Country A, the CPAP survival rate is 85%. In Country B, it's 75%. Naturally, you conclude that Country A's program is more effective. You recommend scaling the Country A model.

But now you disaggregate by severity. Among high-severity cases, Country A's survival rate is 70% and Country B's is 80%. Among low-severity cases, Country A is 90% and Country B is 95%. Country B is better in *both* subgroups. Country A's higher overall rate is entirely an artifact of its patient mix — it treats a much larger share of low-severity cases, which inflates its average.

The aggregated data didn't just fail to identify the better program. It actively recommended the worse one. If you scaled Country A's approach to high-severity patients, you'd get worse outcomes than Country B's approach. The average lied.

Why does this happen? Because aggregation hides confounders. A confounder is a variable that influences both the treatment (which country's program a patient is in) and the outcome (survival). In this case, severity is the confounder. It affects both which program patients end up in (sicker patients disproportionately end up in one program) and the outcome (sicker patients have lower survival regardless of treatment).

When you aggregate across severity levels, you're effectively comparing apples to oranges while convincing yourself they're all fruit. The disaggregated data compares apples to apples, and the answer reverses.

Now, here's the harder question: should you always disaggregate? Is the within-group analysis always right and the aggregate always wrong?

No. And this is where understanding causal structure becomes essential.

If the variable you're considering disaggregating on is a confounder — a common cause of both treatment and outcome — then you *should* disaggregate (or statistically adjust). Failing to do so produces a biased estimate.

But if the variable is a mediator — something that lies on the causal pathway *between* treatment and outcome — then disaggregating is wrong. It removes part of the treatment's effect and gives you a biased estimate in the other direction.

Consider this: a training program improves staff knowledge, which improves care quality, which reduces mortality. Knowledge is a mediator. If you compare trained vs. untrained staff while holding knowledge constant, you'll find that training has no effect — because you've blocked the very pathway through which training works. You've disaggregated on the mechanism and then wondered why the mechanism disappeared.

The difference between a confounder and a mediator is not in the data. They can produce identical statistical signatures. The difference is in the causal model — in the direction of the arrows. A confounder points into both treatment and outcome (a common cause). A mediator points from treatment through to outcome (an intermediate step). Only the causal story tells you which is which.

This is why Simpson's Paradox isn't really a statistical problem. It's a causal problem. The data alone cannot tell you whether to aggregate or disaggregate. You need a causal model to decide.

There's a closely related trap called the Prosecutor's Fallacy, and it shows up in diagnostic testing.

Suppose you have a screening test for preeclampsia that is 95% sensitive (it catches 95% of true cases) and 90% specific (it correctly clears 90% of non-cases). Sounds pretty good. A health administrator might see these numbers and mandate universal screening.

But here's the question: if a woman tests positive, what's the probability she actually has preeclampsia?

The answer depends entirely on the base rate. If preeclampsia prevalence is 5%, then among 1,000 women, 50 have the condition and 950 don't. The test catches 47.5 of the 50 true cases (95% sensitivity). But it also falsely flags 95 of the 950 healthy women (10% false positive rate). So of the 142.5 positive results, only 47.5 are true positives. The probability of actually having preeclampsia given a positive test is only about 33%.

A test that is 95% accurate produces correct results only a third of the time when the condition is rare. This is not a flaw in the test — it's a consequence of base rates. And confusing test accuracy with predictive value is the Prosecutor's Fallacy.

In courtrooms, this fallacy has sent innocent people to prison. ("The probability of this DNA match occurring by chance is one in a million, therefore the defendant is almost certainly guilty." But if you tested a million people, you'd expect one false match — and that match would be indistinguishable from a true match.)

In global health, the Prosecutor's Fallacy leads to misallocation of diagnostic resources, over-treatment of false positives, and false confidence in screening programs. The cure for it is Bayes' Rule: always combine test accuracy with prior probability.

Both Simpson's Paradox and the Prosecutor's Fallacy share a common root: they emerge when you ignore the structure of the problem and focus only on surface-level numbers. Aggregated success rates, test accuracy percentages — these are not wrong numbers. They're incomplete numbers. They answer a different question than the one you're asking, and the gap between those questions can be measured in lives.

The lesson is uncomfortable but essential: data does not speak for itself. It speaks through the causal model you use to interpret it. The same data set, filtered through different causal assumptions, can tell opposite stories. And there is no statistical test that will tell you which story is correct. Only the causal model — the set of arrows you've drawn, the structure of influence you've committed to — can guide you to the right interpretation.

In cost-effectiveness analysis for maternal and newborn health, these issues are not academic. When you compute cost per life saved across countries, you're aggregating across wildly different contexts — different disease burdens, different baseline capacities, different levels of government co-financing. The aggregate number might suggest that investing in Country X is twice as cost-effective as investing in Country Y. But disaggregated by facility type, the ranking might reverse.

The defense against these traps is not more data. It's more structure. Draw the causal model first. Identify the confounders and mediators. Then — and only then — decide how to slice the data.

The average is the most natural statistic in the world. It's also, in the wrong context, the most dangerous one. Trust it only after you've checked what it's hiding.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 7,
    "title": "The Difference Between Watching and Doing",
    "subtitle": "Why giving a facility a CPAP machine is not the same as noticing it already has one",
    "body": r"""
Here is a sentence that sounds like a tautology but is actually one of the deepest ideas in all of causal reasoning:

Seeing that something is true is not the same as making it true.

Let me explain why this matters, and why confusing the two has led to some of the most expensive mistakes in global health.

Imagine you're looking at data from 200 health facilities. You notice that facilities with CPAP machines have significantly better neonatal outcomes than facilities without them. The correlation is strong, it survives controlling for obvious factors like facility size and urban/rural location, and it's consistent across regions. A naive conclusion: CPAP machines improve outcomes. Buy more CPAPs.

But wait. Why do some facilities have CPAPs and others don't? Maybe it's because the facilities with CPAPs are in districts with stronger health management — districts that invest in both equipment and staff training, that attract better clinicians, that have reliable supply chains. The CPAP machine is a marker of a well-functioning system, not necessarily the cause of good outcomes.

When you *observe* that a facility has a CPAP, you're seeing the result of a process. That process selected certain facilities for CPAP placement, and the same factors that drove the selection also drive outcomes. The observation is contaminated by selection.

When you *intervene* to give a facility a CPAP — especially a facility that wouldn't otherwise have one — you're doing something fundamentally different. You're bypassing the selection process. You're breaking the causal link between "the reasons a facility gets a CPAP" and "the CPAP itself." The facility now has a CPAP, but it doesn't have the strong management, trained staff, and reliable supply chain that typically come with it.

This distinction — between observing a variable in its natural state and intervening to set it to a particular value — is formalized in what's called the *do-operator*, and it's the foundation of modern causal inference.

In mathematical notation, P(Outcome | CPAP = yes) is the probability of a good outcome among facilities that *happen to have* a CPAP. P(Outcome | do(CPAP = yes)) is the probability of a good outcome if you *give* a facility a CPAP. These are different quantities. They can have different values. And confusing them is one of the most common and consequential errors in program evaluation.

The mechanics of the do-operator are surprisingly elegant. When you intervene on a variable, you "perform surgery" on the causal graph. You cut all the arrows pointing *into* the intervened variable and replace them with your intervention. The variable is no longer caused by its parents — it's caused by you. But all the arrows pointing *out* of it remain intact. The variable still causes its downstream effects; it's just no longer a consequence of its upstream causes.

Consider a simple model: Government Commitment → CPAP Availability → Care Quality → Neonatal Outcomes, with Government Commitment also directly affecting Staff Training → Care Quality. When you observe CPAP Availability, you're seeing a quantity that's correlated with Government Commitment (because commitment causes CPAP placement). That correlation propagates through the model, making it look like CPAPs have effects they don't actually have.

When you perform do(CPAP = yes), you cut the arrow from Government Commitment to CPAP Availability. CPAP is now set by fiat. It no longer carries information about government commitment. The only effect of the CPAP on outcomes flows through its actual causal pathway — through care quality — uncontaminated by the confounding path through government commitment.

This is called graph surgery, and it transforms a messy observational question into a clean causal one.

Now let's step back and think about why this matters for decision-making in practice.

When a donor organization is deciding where to allocate funds, they need to evaluate interventions — things they will *do*. Should we buy CPAPs? Should we fund training programs? Should we invest in community health worker networks? Each of these is an intervention, a do-operation. It changes the world in a specific way.

But most of the evidence available to inform these decisions is observational. We have data on which facilities have CPAPs and how they perform. We have data on which districts invested in training and what happened afterward. We have cross-country comparisons of health spending and health outcomes. All of this is observation, not intervention.

The do-operator tells you how to bridge the gap — how to estimate the effect of doing something using data from watching things happen. And the bridge depends entirely on the causal model. With the right model, you can identify which confounders need to be accounted for, which mediators should not be adjusted for, and whether the causal effect is even identifiable from the available data.

This brings us to influence diagrams — an extension of causal models that explicitly incorporate decisions and objectives. An influence diagram has the same probabilistic nodes as a Bayesian network, but it adds decision nodes (rectangles representing choices you control) and objective nodes (hexagons representing what you're trying to optimize).

In our CPAP example, the decision node is "Invest in CPAP program: yes or no." The objective node is "Expected lives saved per dollar." The probabilistic nodes include equipment functionality, staff capacity, facility utilization, and neonatal outcomes. The influence diagram shows how your decision flows through the causal model to affect your objective.

Expected value analysis then becomes straightforward, at least in principle. For each possible decision, you compute the expected value of your objective by propagating the intervention through the model. The decision with the highest expected value wins. If the expected lives saved per dollar from CPAP investment exceeds that from training investment, you buy CPAPs.

But expected value is only as good as your model. And this is where the Expected Value of Perfect Information (EVPI) becomes crucial. EVPI tells you: how much better could you do if you knew, with certainty, the true state of an uncertain variable before deciding? If EVPI for a particular variable is high, it means your decision is sensitive to that uncertainty, and you should invest in learning more before acting. If EVPI is low, the uncertainty doesn't matter much for your decision, and you can proceed.

This is a profound idea. It tells you which uncertainties are worth resolving and which are not. In a world of limited research budgets, EVPI is a rational guide to what questions to study next.

The framework of influence diagrams, do-operators, and expected value analysis transforms decision-making from an art into a craft. It doesn't remove judgment — you still need to build the model, specify the probabilities, and define the objective. But it provides a rigorous scaffolding that prevents the most common errors: confusing observation with intervention, ignoring confounders, and failing to account for the value of additional information.

As the economist Charles Manski has argued, the choice between doing nothing and doing something is itself a decision under uncertainty. A formal framework doesn't make the uncertainty go away. It makes the uncertainty visible, quantifiable, and manageable.

And it all starts with that deceptively simple distinction: watching is not doing. Once you internalize this, you'll never look at a regression table the same way again.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 8,
    "title": "Portfolio Thinking for Lives Saved",
    "subtitle": "Why the same logic that diversifies a stock portfolio applies to saving lives across ten countries",
    "body": r"""
In 1952, a young economist named Harry Markowitz published a paper that would eventually win him the Nobel Prize. His insight was elegant and, in retrospect, obvious: you shouldn't evaluate investments in isolation. The risk of a portfolio depends not just on the risk of each individual asset, but on how those assets move relative to each other. Two risky investments that are uncorrelated can, together, produce a portfolio that is less risky than either one alone.

This idea — portfolio theory — transformed finance. And it has something profound to tell us about how to allocate health resources across countries.

Consider the problem facing a global health donor with $500 million to invest in maternal and newborn health across sub-Saharan Africa. There are ten candidate countries, each with different baseline conditions, different government commitment levels, different cost structures, and different levels of uncertainty. The naive approach is to rank countries by expected cost-effectiveness and pour all the money into the top-ranked country.

This is the equivalent of putting your entire retirement savings into the single stock with the highest expected return. It maximizes expected performance. It also maximizes the probability of catastrophic failure.

Here's why the naive approach fails. Expected cost-effectiveness is estimated with enormous uncertainty. Your model says Country A will save lives at $2,000 each and Country B at $3,000 each. But the confidence intervals overlap massively. Country A's cost might actually be $1,500 — or $8,000. You don't know whether the government will co-finance as promised. You don't know if the supply chain will hold. You don't know if community health workers will be absorbed into the health system or left to drift.

If you put everything into Country A and the government doesn't co-finance, or the supply chain collapses, or there's a political crisis — you've lost not just money, but the lives you were trying to save. And you've lost them completely, with no backup.

Portfolio thinking says: spread across countries, but not equally. Invest more in countries with higher expected returns, but also invest in countries that are uncorrelated — countries where the risks are driven by different factors. If Country A's risk is mainly political and Country B's risk is mainly logistical, a crisis in one is unlikely to coincide with a crisis in the other. Together, they smooth out the overall risk.

This isn't a metaphor. It's the same mathematics. Expected lives saved replaces expected financial return. Uncertainty in cost-effectiveness replaces investment risk. The correlation between country outcomes replaces the correlation between stock returns. And the optimization objective — maximize expected lives saved while keeping downside risk manageable — is directly analogous to the efficient frontier in finance.

But the health allocation problem is harder than a stock portfolio in several important ways.

First, the uncertainties are deeper. In finance, you have decades of historical returns to estimate risk. In global health, you might have one or two similar programs in similar countries, and even those aren't truly comparable. Your probability distributions are based on expert judgment as much as data.

Second, the investments are not liquid. If a stock drops, you can sell it. If a country program is underperforming, you can't easily redirect the funds. Facilities have been built, staff have been hired, community expectations have been set. Exiting a country is costly, slow, and politically fraught. This means your initial allocation matters far more than in finance, where you can rebalance quarterly.

Third — and this is the big one — the countries respond to your investment. Stocks don't know you own them. Governments do. And their behavior changes based on what you invest.

This is the crowding-out problem. When a donor invests heavily in a country's health system, the government may reduce its own spending. "The donors are handling it." The net effect of your investment is less than what your models predicted, because the government took your money and redirected its budget elsewhere. Your $50 million investment might only produce $30 million of additional impact, because the government quietly pulled $20 million from health and spent it on roads.

The opposite can also happen — crowding in. A large donor commitment can signal confidence in the country's health system, making it easier for the government to justify increased health spending to its legislature. "Even the donors believe in our plan." In this case, your $50 million catalyzes an additional $20 million from the government, and total impact exceeds your investment.

Whether you get crowding out or crowding in depends on the government's incentives, its fiscal constraints, its political dynamics, and the design of your investment contract. And it varies by country. This is where game theory meets portfolio theory — a theme we'll explore fully in the next chapter.

For now, let's focus on how to handle the uncertainty itself. One powerful tool is Monte Carlo simulation. Instead of computing a single expected cost-effectiveness for each country, you run thousands of scenarios. In each scenario, you randomly draw values for each uncertain variable from its probability distribution — government co-financing rates, equipment functionality, staff retention, community uptake — and compute the resulting cost-effectiveness. After thousands of draws, you have a distribution of possible outcomes for each country and each allocation strategy.

This lets you answer questions that expected-value analysis cannot. What's the probability that our portfolio saves fewer than 100,000 lives? What's the worst-case scenario? What allocation minimizes the probability of catastrophic underperformance? These are risk management questions, and they require distributional thinking, not point estimates.

Consider three allocation strategies across three countries. Strategy One: concentrate everything in the country with the highest expected return. Strategy Two: spread equally across all three. Strategy Three: a tiered approach — invest most in the highest-expected-return country, but maintain meaningful investments in the other two as hedges.

Monte Carlo simulation might show that Strategy One has the highest expected lives saved but also the widest distribution — including scenarios where it saves very few lives. Strategy Two has a lower expected value but a much tighter distribution — less upside, but much less downside. Strategy Three might hit a sweet spot: nearly as high an expected value as concentration, with nearly as narrow a distribution as diversification.

The "right" answer depends on your risk tolerance. If you're a small donor making a single bet, you might prefer concentration — swing for the fences. If you're a large donor with accountability to multiple stakeholders, you probably prefer the tiered approach — strong expected performance with meaningful protection against downside scenarios.

There's another dimension to this problem that doesn't exist in finance: sequential decision-making. In a stock portfolio, you make decisions at regular intervals. In health allocation, the sequencing matters enormously.

You might commit initial funding to all ten countries, then observe first-year results, then scale up in the countries that are performing well and hold steady (or exit) in the countries that aren't. This "commit-observe-scale" approach uses early investments as information-gathering exercises. You're buying options, not just assets.

The value of this sequential approach depends on how informative the early evidence is. If first-year results are highly predictive of long-term outcomes, the option value is high — you can redirect resources toward winners early. If first-year results are noisy and unreliable, the option value is low — you can't learn much from early data, so you might as well commit upfront.

This is another instance of Expected Value of Information — the same concept we met in influence diagrams, now applied at the portfolio level. How much is it worth to wait and learn before committing fully? The answer is context-specific and quantifiable.

The deeper lesson of portfolio thinking is philosophical. It says: acknowledge that you don't know which country will perform best. Acknowledge that your models are uncertain. And then make that uncertainty work *for* you, by constructing a portfolio that is robust to being wrong about any single country.

In finance, the investors who survive longest are not the ones who make the biggest bets. They're the ones who manage risk intelligently while maintaining exposure to upside. The same is true in global health. The donors who save the most lives over the long term are not the ones who bet everything on their best guess. They're the ones who build portfolios that perform well across a wide range of possible futures.

Diversification isn't about admitting defeat. It's about acknowledging reality.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 9
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 9,
    "title": "When the World Pushes Back",
    "subtitle": "Why governments free-ride on donor investments — and how to design contracts that make cooperation rational",
    "body": r"""
In the early 2000s, a major donor invested $200 million in health systems across several African countries. The investment came with a promise from each government: "We will maintain our current health spending and add your funds on top." The math was simple. Donor money plus government money equals more health spending equals more lives saved.

Within two years, three of the five governments had quietly reduced their health budgets. Not dramatically — just enough to redirect funds to other priorities. Roads. Security. Election-year projects. The donor's $200 million didn't produce $200 million in additional health spending. It produced maybe $120 million, with governments pocketing the difference.

The donor was outraged. The governments were rational.

This is the fundamental insight of game theory: the other player has a strategy too. And until you account for that strategy — until you understand what's rational for them, not just what's ethical or what they promised — your plans are built on sand.

The health funding landscape is, at its core, a strategic interaction between donors and governments. Each side has objectives, constraints, and incentives. Each side's optimal behavior depends on what the other side does. And the outcome depends on both strategies simultaneously.

Let's formalize this with the simplest possible game. A donor and a government each decide how much to invest in health. The donor can invest High or Low. The government can invest High or Low. The total health impact depends on the combined investment.

If both invest High, the country gets excellent health outcomes. If both invest Low, outcomes are poor. If the donor invests High and the government invests Low, outcomes are moderate — better than nothing, because the donor's money still buys some services, but worse than the cooperative outcome because the government's complementary investments (salaries, supply chains, management) are missing.

Now here's the key: what's the government's best response to each donor strategy?

If the donor invests High, the government's best move is to invest Low. Why? Because the donor is already covering much of the cost. The government gets moderate outcomes with low spending — a decent deal. Investing High would improve outcomes only marginally while costing the government a lot.

If the donor invests Low, the government's best move is... also to invest Low. The donor isn't contributing much, so the government would have to bear most of the cost alone for only marginally better outcomes.

In both cases, the government's best response is to invest Low. And if the donor reasons symmetrically — knowing the government will invest Low regardless — the donor's best response is also to invest Low.

This is the Nash Equilibrium: both sides investing Low, even though both would be better off if both invested High. It's the Prisoner's Dilemma in a health funding context. Each side, acting rationally in its own interest, produces an outcome that neither side wants.

The technical name for this is the free-rider problem. Each player has an incentive to let the other bear the cost. The result is underinvestment by everyone.

If this sounds abstract, consider the Workforce Absorption Game — a real and pressing problem in global health.

A donor funds training for 500 community health workers (CHWs). The program costs $5 million over three years. The expectation — the hope — is that the government will absorb these workers into the formal health system when the donor funding ends. The workers need salaries, supervision, supplies, and integration into the district health management structure.

The government, however, faces a different calculus. Absorbing 500 CHWs costs money — not just salaries, but the entire support system. And the government has other priorities. Its leaders know that if the donor cares enough about these CHWs, the donor will extend funding rather than let the program collapse. After all, the donor has already invested $5 million. Walking away means admitting failure. The sunk cost works in the government's favor.

So the government delays. It promises absorption "next fiscal year." It establishes a committee to study the integration process. It cites budget constraints. And the donor, faced with the choice between extending funding or watching 500 trained health workers drift away, extends.

The CHWs are now perpetually donor-funded. The government has successfully free-ridden on the donor's commitment. And the donor is locked into a recurring cost it never intended to bear.

How do you escape this trap?

Game theory offers several mechanisms, all based on the same principle: change the structure of the game so that cooperation becomes the rational strategy.

The first mechanism is commitment devices. Make free-riding costly. For example, structure the funding as milestone-based disbursements. The donor commits $5 million, but releases funds in tranches tied to government actions. Tranche 1: released when the government publishes an integration policy. Tranche 2: released when CHW positions appear in the government budget. Tranche 3: released when the first salary payments are made from government funds.

This changes the government's calculus. Free-riding no longer just means getting the donor to keep paying. It means losing the next tranche. The cost of non-cooperation is concrete and immediate.

The second mechanism is repeated interaction. In a one-shot game, free-riding is rational because there are no future consequences. But donors and governments interact repeatedly — across years, across programs, across sectors. In a repeated game, a government that free-rides today risks losing donor confidence (and donor funding) tomorrow. The shadow of the future creates an incentive for cooperation.

This is why reputation matters enormously in international development. A government that has a track record of co-financing, of absorbing programs, of meeting commitments, attracts more and better donor partnerships. A government that consistently free-rides finds itself gradually isolated — offered smaller grants with stricter conditions. Reputation is the currency of repeated games.

The third mechanism is game design — structuring the interaction so that the cooperative outcome is also the individually rational outcome. This is harder, but it's the gold standard.

Consider a matching grant structure. The donor matches every dollar the government spends on health up to a ceiling. Now the government's incentive changes completely. Every dollar it spends buys two dollars of health services. The cost of investing High is halved. Free-riding — spending Low — means leaving matching funds on the table. The government's rational response to a matching grant is to invest more, not less.

This is not hypothetical. Gavi, the Vaccine Alliance, uses a co-financing mechanism that gradually shifts costs from donors to governments as country income rises. Countries start by paying 20 cents per dose, with Gavi covering the rest. Over time, the country share increases. By the time a country "graduates" from Gavi support, it's paying the full cost — and the transition is gradual enough that government budgets can adjust.

The design works because it aligns incentives rather than relying on promises. The government doesn't co-finance because it promised to. It co-finances because co-financing is the rational strategy given the game structure.

This is the central lesson of game theory for global health: don't rely on good intentions. Design the game so that rational self-interest leads to cooperative outcomes.

There's a subtlety here that's easy to miss. The point of game theory is not to be cynical about governments. Most health ministers genuinely want to improve outcomes for their populations. The problem is that they operate within systems — political systems, budgetary systems, bureaucratic systems — that create incentives. And those incentives sometimes point away from cooperation, even when the individuals involved want to cooperate.

A health minister who knows that donor funding will continue regardless of government spending faces a genuine dilemma. She could fight for a larger health budget, spending political capital that she needs for other battles. Or she could accept the donor funding, maintain the status quo, and use her political capital elsewhere. The second option is rational, even if it's not ideal for health outcomes.

Game theory doesn't blame her for this choice. It acknowledges it, analyzes it, and then asks: how can we change the structure so that fighting for the health budget becomes the rational choice?

The most effective health investments in the coming decades will not be the ones with the best technical design. They'll be the ones with the best strategic design — the ones that anticipate how governments, communities, and health workers will respond, and that channel those responses toward cooperation.

Because in a world of strategic actors, the plan that ignores everyone else's strategy isn't a plan at all. It's a wish.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 10
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 10,
    "title": "Can Data Discover Causes?",
    "subtitle": "Why you need both algorithms and domain expertise — and why neither alone is enough",
    "body": r"""
There's a fantasy that haunts the data science community. It goes like this: give me enough data and a powerful enough algorithm, and I can discover the causal structure of any system — without needing to understand the system at all. Feed hospital records into a machine, and it will tell you what causes disease. Feed economic data into a model, and it will tell you what drives growth. Feed health facility data into a pipeline, and it will tell you which investments save lives.

This fantasy is not entirely wrong. And that's what makes it so dangerous.

Over the past two decades, a set of algorithms has been developed that can, under certain conditions, learn causal structure from observational data. These algorithms — with names like the PC algorithm, GES, and hill-climbing search — analyze patterns of conditional independence in data to infer which variables cause which. They're mathematically rigorous, computationally tractable, and genuinely impressive.

They're also fundamentally limited in ways that are easy to overlook.

Let me explain both the promise and the limitation, because understanding the boundary between what data can and cannot tell you about causation is one of the most important skills you'll develop in this course.

Start with the promise. Two variables, A and B, are correlated. There are three possible causal explanations: A causes B, B causes A, or some third variable C causes both. In many cases, you can distinguish between these by looking at conditional independence patterns involving other variables.

Here's the intuition. If A causes B causes C (a chain), then A and C are correlated, but they become independent when you condition on B. If A and C are both caused by B (a fork), you get the same pattern — correlation between A and C that vanishes when you condition on B. But if A and C both cause B (a collider), the pattern reverses: A and C are *independent*, but they become *dependent* when you condition on B.

This asymmetry between colliders and non-colliders is the key that unlocks causal discovery. By systematically testing conditional independence relationships across all triples of variables, algorithms can identify which structures are consistent with the data and which are not.

The PC algorithm, named after its creators Peter Spirtes and Clark Glymour, does exactly this. It starts with a complete graph (every variable connected to every other) and systematically removes edges that correspond to conditional independences. Then it orients edges based on collider detection. The result is a partially oriented causal graph — a picture of the causal structure that's consistent with the observed data.

"Partially oriented" is the key qualifier. And this is where the limitation begins.

Consider three variables: A — B — C in a chain. The conditional independence patterns are identical regardless of whether A causes B causes C, or C causes B causes A, or B causes both A and C (a fork). These three structures are "Markov equivalent" — they encode the same set of conditional independences and therefore cannot be distinguished by any algorithm using observational data alone.

This is not a limitation of current algorithms. It's a mathematical theorem. No algorithm, no matter how sophisticated, can distinguish between Markov equivalent structures from observational data. The data simply doesn't contain the information needed to tell these structures apart.

What the data *can* identify is which variables are not in the same equivalence class — which edges are definitively oriented. Specifically, it can identify colliders (A → B ← C), because colliders have unique independence patterns. But for non-collider structures, the direction of causation remains ambiguous.

In practice, this means that data-driven causal discovery gives you a partial answer. It tells you which variables are causally related, and it orients some edges (the ones involved in collider structures). But for many edges, you're left with "these variables are related, but we can't tell which direction from data alone."

This is where domain expertise enters — not as a crutch, but as an essential complement.

A health system expert looking at an edge between "staff training" and "care quality" immediately knows the direction: training causes quality, not the other way around. This isn't in the data. It's in the understanding of how the world works. The expert brings temporal knowledge (training happens before its effects), mechanistic knowledge (trained staff apply better protocols), and institutional knowledge (training programs are inputs, quality metrics are outputs).

The most powerful approach combines both: let the algorithm discover the skeleton of the graph (which variables are related), then use domain expertise to orient the ambiguous edges. Neither approach alone is sufficient. The algorithm without expertise produces ambiguous graphs. The expertise without the algorithm is subject to confirmation bias, incomplete mental models, and the inability to process complex conditional independence patterns across dozens of variables.

There's another limitation of data-driven discovery that practitioners need to understand: the faithfulness assumption. Causal discovery algorithms assume that every conditional independence in the data corresponds to a structural feature of the causal graph (specifically, d-separation). This assumption fails when two causal pathways between variables happen to exactly cancel each other out, producing a conditional independence that doesn't correspond to any separation in the graph.

This sounds pathological, and in many applications it is rare. But in engineered systems and policy environments — where interventions are designed to achieve specific targets — exact cancellation is more plausible. If a government adjusts its spending to exactly offset donor contributions (perfect crowding out), the data might show no association between donor spending and total health spending, even though a causal link exists. The faithfulness assumption breaks, and the algorithm draws the wrong conclusion.

Beyond algorithmic discovery, there are other data-driven approaches to causal inference that deserve mention. Instrumental variables (IV) analysis, for instance, exploits a source of variation that affects the treatment but is otherwise unrelated to the outcome. If rainfall affects road conditions, and road conditions affect whether a mother delivers at a facility, and rainfall has no direct effect on neonatal outcomes, then rainfall is an "instrument" for facility delivery. You can use the variation in delivery rates caused by rainfall — and only that variation — to estimate the causal effect of facility delivery on outcomes.

The elegance of IV analysis is that it doesn't require you to measure and adjust for all confounders. It only requires you to find a valid instrument — a variable that satisfies the exclusion restriction (it only affects the outcome through the treatment) and the relevance condition (it actually affects the treatment).

The danger of IV analysis is that the exclusion restriction is untestable. You have to argue, on substantive grounds, that the instrument doesn't affect the outcome through any channel other than the treatment. And these arguments are often more fragile than they appear. Does rainfall really have no direct effect on neonatal outcomes? What about its effect on maternal nutrition, on disease exposure, on health worker attendance? If any of these channels exist, the instrument is invalid and the causal estimate is biased.

This brings us full circle. Whether you're using algorithmic causal discovery, instrumental variables, or good old-fashioned regression analysis, the quality of your causal conclusions depends on the quality of your causal assumptions. Data can refine, test, and occasionally overturn those assumptions. But it cannot generate them from nothing.

The correct workflow — the one that this entire course has been building toward — looks like this:

First, build a qualitative causal model based on domain expertise and stakeholder interviews. Draw the nodes. Draw the arrows. Identify what you believe and where you're uncertain.

Second, quantify the model with probability distributions and conditional probability tables. Put numbers on your beliefs.

Third, use the model for situational analysis, intervention design, and resource allocation. Make decisions.

Fourth, use data to test and refine the model. Run causal discovery algorithms to check whether the data is consistent with your structure. Use Bayesian updating to refine your probability estimates. Use instrumental variables or natural experiments to estimate specific causal effects.

Fifth, update the model and iterate. The model evolves as evidence accumulates.

This workflow treats expert knowledge and data as complementary inputs to a single analytical framework. The expert brings structure; the data brings calibration. The expert proposes; the data disposes — or confirms, or refines.

The scientists who built the field of causal inference — Judea Pearl, Peter Spirtes, Jamie Robins, and others — all converged on the same conclusion: causation is not in the data. Causation is in the model. Data can inform the model, test the model, and update the model. But the model — the set of arrows and nodes that represents your theory of how the world works — must come first.

This might sound like a limitation. It is. But it's also a liberation. It means that causal analysis is not a passive exercise in data processing. It's an active exercise in thinking — in drawing, questioning, debating, and refining your understanding of the system you're trying to change.

The data doesn't speak for itself. It never did. But with the right model, it can speak volumes.
""",
})


# ── Document Generation ───────────────────────────────────────────────────────

def build_document(essays, output_path):
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # -- Heading styles --
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)  # BKA primary blue
        if level == 1:
            h.font.size = Pt(26)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(4)
        elif level == 2:
            h.font.size = Pt(16)
            h.font.italic = True
            h.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(18)
        elif level == 3:
            h.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # ── Title Page ─────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Inference and Intervention")
    run.font.name = "Calibri Light"
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
    run.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("Pre-Class Essays")
    run.font.name = "Calibri Light"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("BK Advisors")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)

    doc.add_page_break()

    # ── Table of Contents (manual) ─────────────────────────────────────────
    toc_heading = doc.add_heading("Contents", level=1)

    for essay in essays:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.space_after = Pt(4)
        if essay["number"] == 0:
            toc_label = f"Introduction:  {essay['title']}"
        else:
            toc_label = f"Chapter {essay['number']}:  {essay['title']}"
        run = toc_para.add_run(toc_label)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # Subtitle line
        sub_para = doc.add_paragraph()
        sub_para.paragraph_format.space_after = Pt(10)
        sub_para.paragraph_format.left_indent = Cm(1.27)
        run = sub_para.add_run(essay["subtitle"])
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.italic = True

    doc.add_page_break()

    # ── Essays ─────────────────────────────────────────────────────────────
    for i, essay in enumerate(essays):
        # Chapter heading
        if essay["number"] == 0:
            doc.add_heading("Introduction", level=3)
        else:
            doc.add_heading(f"Chapter {essay['number']}", level=3)
        doc.add_heading(essay["title"], level=1)
        doc.add_heading(essay["subtitle"], level=2)

        # Body text — split into paragraphs
        body = essay["body"].strip()
        paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]

        for para_text in paragraphs:
            # Handle line continuations within paragraphs
            clean_text = " ".join(para_text.split("\n"))
            p = doc.add_paragraph(clean_text)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(8)

        # Page break after each essay except the last
        if i < len(essays) - 1:
            doc.add_page_break()

    # ── Save ───────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Total essays: {len(essays)}")


if __name__ == "__main__":
    output = os.path.join(os.path.dirname(__file__), "pre-class-essays.docx")
    build_document(essays, output)
