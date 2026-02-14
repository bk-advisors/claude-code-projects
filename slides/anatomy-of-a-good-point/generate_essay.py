"""
Generate Pre-Class Essay for "The Anatomy of a Good Point"
Output: pre-class-essay.docx
Style: Morgan Housel (conversational, story-driven, counterintuitive insights)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ── Essay Content ──────────────────────────────────────────────────────────────

essay = {
    "title": "The Idea That Won the Room",
    "subtitle": "Why structure beats substance in the art of persuasion",
    "body": r"""
There were two people in a conference room in Lilongwe, and only one of them changed anything.

The first was a health economist. She had spent three months analyzing Malawi's nurse-midwife workforce crisis. She had the numbers cold: 679,000 annual births, 22 obstetricians in the entire country, a 65% vacancy rate for nurse-midwife technicians in high-burden facilities. She had modeled the mortality impact. She had cross-referenced global evidence. She knew, with genuine confidence, exactly what needed to happen.

She stood up and presented for twelve minutes. She started with the methodology. She explained the data sources. She walked through three sensitivity analyses. She showed a regression table. And then, near the end, she said something like, "So based on all of this, it seems like we should probably prioritize recruitment of about 200 additional nurse-midwives." Polite nods. A few questions about the methodology. The meeting moved on.

The second person was a program director. He had maybe a quarter of the data. He stood up and said: "We need to recruit 200 nurse-midwives in Malawi's high-burden facilities by year-end. Here's why, and here's what happens if we don't." He took ninety seconds. The room leaned in. By the end of the week, the recommendation was in the donor's investment memo.

Same room. Same recommendation. Completely different outcomes. And the difference had nothing to do with who was smarter or who had better data. The difference was architecture.

This is the story of how arguments actually work — and why the smartest person in the room often isn't the most persuasive one.

Most of us were trained to build arguments the way we build analysis: bottom-up. Start with the data. Build the case brick by brick. Save the conclusion for the end, like the climax of a well-constructed paper. This feels rigorous. It feels honest. It feels like you're earning your conclusion.

But here's the problem: your audience isn't reading a paper. They're sitting in a meeting, or scanning an email, or half-listening on a call. They don't have the patience — or the context — to follow your logic from the ground up. By the time you reach your conclusion, they've already formed their own, and it might not be the one you intended.

The health economist did nothing wrong analytically. Her analysis was sound. Her recommendation was correct. But she committed the most common structural error in professional communication: she delivered her argument in the order she *thought* about the problem, not in the order her audience could *process* it.

There's a framework for fixing this, and it's almost embarrassingly simple. Four steps. You can learn them in five minutes and use them for the rest of your career.

Step one: state the point. Not the context. Not the background. The point. One sentence, up front, before anything else. "We need to recruit 200 nurse-midwives by year-end." That's a point. "I've been looking at the staffing data and there are some interesting patterns" is not. The test is simple: if your opening sentence could be a newspaper headline, you're on track. If it couldn't, you're burying the lead.

Step two: explain why it's true. This is the reasoning — the logical backbone that connects your claim to reality. Not evidence yet. Just logic. "Because Malawi has one neonatologist and a 65% nurse-midwife vacancy rate for 679,000 annual births, and you cannot deliver quality emergency obstetric care with those numbers." Notice what this does. It doesn't cite a study. It doesn't tell a story. It lays out a causal chain: here are the facts, here's why they're insufficient, here's the conclusion. That's the "because."

Step three is where most people do half the work. You need evidence, but you need it in two directions.

First, zoom in. Pick one specific case and make it vivid. "In one district hospital in Malawi, there were zero trained neonatal nurses covering a unit with 30 admissions per day. A partner program placed just four mentored nurse-midwives there last year. Neonatal mortality dropped 35% in six months." That's a concrete, tangible story. People remember it. It makes the abstract real.

But if you stop there, someone in the room will think: that's one hospital. Maybe it was an outlier. Maybe the timing was coincidental. So you need the second move: zoom out. "This isn't just one facility. The Lancet's modeling shows that a 10% increase in midwifery coverage every five years could avert 22% of maternal deaths, 23% of newborn deaths, and 14% of stillbirths globally — saving 1.35 million lives by 2035." That's the pattern. It tells the audience: this works everywhere, not just in that one district.

Zoom in makes you *feel* the point. Zoom out makes you *believe* it. You need both. An anecdote without data is a story. Data without an anecdote is a spreadsheet. Neither, on its own, is persuasive.

Step four — and this is the one almost everyone forgets — answer the question your audience is silently asking the entire time: "So what?"

"If we fill these 200 positions, we can prevent over 3,000 maternal and newborn deaths per year in Malawi alone. If we don't, those positions stay vacant and those deaths remain preventable but unprevented. Every month of delay costs lives we have the tools to save."

That's the impact. That's what turns a well-constructed case into a reason to act. Without step four, you've made an intellectually sound argument. With it, you've given people something they can't ignore.

Now here's the counterintuitive part. You might think the key to persuasion is having better data, or deeper expertise, or more polished slides. It's not. The key is completeness.

A complete argument with average data beats a brilliant argument with missing pieces. Every time. Because people don't reject ideas because the evidence is weak. They reject ideas because the argument feels incomplete — it leaves questions unanswered. "That's interesting, but why should I care?" "That's one example, but is it a pattern?" "That's a pattern, but what does it look like on the ground?" Each missing step is a gap that doubt rushes in to fill.

The framework isn't magic. It's a checklist against incompleteness. Point. Reasoning. Zoom In. Zoom Out. Impact. If any box is empty, you're not ready. If all four are filled, you probably can't be ignored.

Think back to the program director in that conference room. He didn't have three months of analysis. He didn't have sensitivity tables. But he had all four pieces: a clear claim, a logical reason, a specific example, a global pattern, and a cost of inaction. His argument was complete. The economist's was thorough but incomplete — she had buried the point and forgotten the "so what."

There's a phrase worth remembering: a good point isn't just correct — it's complete.

The next time you need to make an argument — in a meeting, in an email, in a presentation — don't start by collecting more data. Start by filling in four boxes. What's the claim? What's the logic? What's one specific example? What's the bigger pattern? And what happens if we don't act?

If you can answer all four in under sixty seconds, you're ready. And if you can't, now you know exactly which piece is missing.
""",
}

# ── Document Builder ───────────────────────────────────────────────────────────


def build_document(essay, output_path):
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # -- Default paragraph font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # -- Heading styles --
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)  # BKA primary blue
        if level == 1:
            h.font.size = Pt(26)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(4)
        elif level == 2:
            h.font.size = Pt(16)
            h.font.italic = True
            h.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(18)
        elif level == 3:
            h.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # ── Title Page ─────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("The Anatomy of a Good Point")
    run.font.name = "Calibri Light"
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
    run.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("Pre-Class Essay")
    run.font.name = "Calibri Light"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("BK Advisors")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)

    doc.add_page_break()

    # ── Essay ──────────────────────────────────────────────────────────────
    doc.add_heading(essay["title"], level=1)
    doc.add_heading(essay["subtitle"], level=2)

    # Body text — split into paragraphs
    body = essay["body"].strip()
    paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]

    for para_text in paragraphs:
        # Handle line continuations within paragraphs
        clean_text = " ".join(para_text.split("\n"))
        p = doc.add_paragraph(clean_text)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)

    # ── Save ───────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    output = os.path.join(os.path.dirname(__file__), "pre-class-essay.docx")
    build_document(essay, output)
