"""
Generate YouTube Upload Metadata for Storytelling with Data Videos
Output: youtube-metadata-intro-video.docx, youtube-metadata-full-lecture.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# -- Metadata Content -------------------------------------------------------

intro_video = {
    "title": "Storytelling with Data: The Missing Pieces Most Courses Skip | Lecture Intro",
    "description": (
        "Most data storytelling frameworks teach you how to make a chart. "
        "But they skip two critical thinking modes that turn charts into arguments.\n\n"
        "In this 5-minute introduction, Matthew Kuch shares what a decade of data "
        "visualization practice — from simple bar charts in 2015 to interactive "
        "scrollytelling in 2026 — taught him about the diagnostic and predictive "
        "layers that most data stories miss.\n\n"
        "This intro previews the full lecture on Storytelling with Data, "
        "part of BK Advisors' professional skills series for global health consultants.\n\n"
        "YOU'LL LEARN:\n"
        "- Why curiosity (not data) is the starting point for every data story\n"
        "- The three-lens framework: Diagnostic, Descriptive, Predictive\n"
        "- The five-step process for making a data story\n"
        "- The diagnostic lens: asking 'why did this happen?'\n"
        "- The predictive lens: asking 'what comes next?'\n"
        "- How to complete the full scientific method loop in data storytelling\n\n"
        "CASE STUDIES:\n"
        "- Africa Measles Heatmap (R/Quarto): https://bk-advisors.github.io/africa-measles/\n"
        "- HPV Vaccination in PNG (Svelte): https://bk-advisors.github.io/hpv-png-story/\n\n"
        "ABOUT THE SPEAKER:\n"
        "Matthew Kuch has spent over a decade in global child health — at CHAI in Uganda, "
        "Gavi in Geneva managing country portfolios across Kenya and Lesotho, and at Deloitte "
        "managing grants for USAID and DFID. He has been practicing data visualization since "
        "2015 and builds interactive data stories for public health advocacy.\n\n"
        "BK Advisors | Professional Development for Global Health Consultants"
    ),
    "tags": (
        "data storytelling, data visualization, dataviz, storytelling with data, "
        "global health consulting, maternal and newborn health, MNH, "
        "data analysis, consulting skills, professional development, "
        "BK Advisors, Matthew Kuch, measles vaccination, HPV vaccination, "
        "heatmap, scrollytelling, interactive visualization, "
        "diagnostic analysis, predictive analysis, scientific method, "
        "Morgan Housel, lowercase storytelling, R, ggplot2, Svelte, Quarto"
    ),
    "category": "Education",
    "thumbnail_suggestions": (
        "TEXT OVERLAY: \"The Missing Pieces\" in large white bold text, "
        "with \"in Data Storytelling\" below in yellow (#FED141)\n\n"
        "VISUAL: Four-quadrant diagram showing Diagnostic/Descriptive/"
        "Prescriptive/Predictive, with the Diagnostic and Predictive "
        "quadrants highlighted and the other two faded. "
        "Use BK Advisors primary blue (#005CB9) background.\n\n"
        "ALTERNATIVE: Split screen — left side shows a standard bar chart "
        "(labeled 'What most people do'), right side shows the four-quadrant "
        "model (labeled 'The full loop')."
    ),
    "chapters": (
        "0:00 — The Chart That Started a Project\n"
        "0:30 — About the Speaker\n"
        "1:00 — What This Session Covers\n"
        "1:30 — The Three-Lens Framework\n"
        "2:30 — The Five-Step Process (Preview)\n"
        "3:30 — The Missing Pieces: Diagnostic & Predictive\n"
        "4:30 — Case Studies Overview\n"
        "5:00 — Your Challenge"
    ),
}

full_lecture = {
    "title": "Storytelling with Data: Completing the Scientific Method Loop | Full Lecture — BK Advisors",
    "description": (
        "Every data story has a finding. The best data stories also have a diagnosis "
        "and a prediction. This lecture teaches you how to build all three.\n\n"
        "Using two real case studies from global health — an interactive measles heatmap "
        "covering 43 years across Africa, and a scrollytelling experience making the case "
        "for HPV vaccination in Papua New Guinea — Matthew Kuch walks through a complete "
        "framework for data storytelling that goes beyond charts and dashboards.\n\n"
        "WHAT YOU'LL LEARN:\n"
        "- Where data stories start (hint: not with data)\n"
        "- Why the standard analysis pipeline misses what matters\n"
        "- How to start with an audience of one\n"
        "- The five-step making process: Question, Measure, Analyze, Synthesize, Form\n"
        "- The exploration vs. explanation distinction\n"
        "- The diagnostic lens: identifying causal mechanisms in your data\n"
        "- The predictive lens: projecting findings forward in time\n"
        "- The four-quadrant model: Diagnostic, Descriptive, Prescriptive, Predictive\n"
        "- How to complete the full scientific method loop\n"
        "- The compound effect of lowercase storytelling\n\n"
        "CASE STUDIES:\n"
        "- Africa Measles Heatmap: https://bk-advisors.github.io/africa-measles/\n"
        "- HPV Vaccination in PNG: https://bk-advisors.github.io/hpv-png-story/\n\n"
        "INSPIRED BY:\n"
        "Aman Bhargava's 'Data Storytelling in Lowercase' guest talk at Purdue University\n"
        "https://aman.bh/blog/2026/data-storytelling-in-lowercase-guest-talk-at-purdue\n\n"
        "SERIES: Professional Skills for Global Health Consultants\n"
        "LECTURE 1: The Anatomy of a Good Point\n"
        "LECTURE 2: Making Causal Arguments That Hold Up\n"
        "THIS VIDEO: Storytelling with Data\n\n"
        "ABOUT THE SPEAKER:\n"
        "Matthew Kuch has spent over a decade in global child health — at CHAI in Uganda, "
        "Gavi in Geneva managing country portfolios across Kenya and Lesotho, and at Deloitte "
        "managing grants for USAID and DFID. He has been practicing data visualization since "
        "2015 and builds interactive data stories for public health advocacy.\n\n"
        "BK Advisors | Professional Development for Global Health Consultants"
    ),
    "tags": (
        "data storytelling, data visualization, dataviz, storytelling with data, "
        "global health consulting, maternal and newborn health, MNH, "
        "data analysis, consulting skills, professional development, "
        "BK Advisors, Matthew Kuch, measles vaccination, HPV vaccination, "
        "heatmap, scrollytelling, interactive visualization, "
        "diagnostic analysis, predictive analysis, scientific method, "
        "Morgan Housel, lowercase storytelling, R, ggplot2, Svelte, Quarto, "
        "four quadrant model, exploration vs explanation, data journalism, "
        "public health data, Africa measles, Papua New Guinea, cervical cancer, "
        "WHO UNICEF measles initiative, evidence-based policy"
    ),
    "category": "Education",
    "thumbnail_suggestions": (
        "TEXT OVERLAY: \"Storytelling with Data\" in white bold text, "
        "with \"FULL LECTURE\" badge in BK Advisors green (#83BD00)\n\n"
        "VISUAL: The four-quadrant model — Diagnostic (orange), Descriptive (light blue), "
        "Prescriptive (green), Predictive (teal) — with axis labels 'Past/Future' and "
        "'Understanding/Action'. Use BK Advisors primary blue (#005CB9) background.\n\n"
        "ALTERNATIVE: Two case study screenshots side by side — the Africa measles heatmap "
        "on the left and the HPV-PNG scrollytelling on the right — with text overlay "
        "\"From Diagnosis to Prediction\" and the BK Advisors logo."
    ),
    "chapters": (
        "NOTE: Timestamps below are approximate based on lecture structure.\n"
        "Adjust to match the actual video after recording.\n\n"
        "0:00 — Introduction: The Chart That Started a Project\n"
        "XX:XX — Part I: Philosophy\n"
        "XX:XX — Where It Starts: Curiosity\n"
        "XX:XX — What's the Norm: The Linear Process\n"
        "XX:XX — Start with an Audience of One\n"
        "XX:XX — Part II: Making a Story\n"
        "XX:XX — Step 1: Find Your Question\n"
        "XX:XX — Step 2: Make It Measurable\n"
        "XX:XX — Step 3: Analyze\n"
        "XX:XX — Step 4: Synthesize\n"
        "XX:XX — Step 5: Choose Your Form\n"
        "XX:XX — Part III: The Missing Pieces\n"
        "XX:XX — The Diagnostic Lens\n"
        "XX:XX — Case Study: Africa Measles Heatmap\n"
        "XX:XX — The Predictive Lens\n"
        "XX:XX — Case Study: HPV Vaccination in PNG\n"
        "XX:XX — Completing the Loop: Four Modes of Data Storytelling\n"
        "XX:XX — The Full Loop\n"
        "XX:XX — The Compound Effect\n"
        "XX:XX — Closing"
    ),
}

# -- Document Builder -------------------------------------------------------


def build_document(metadata, video_filename, output_path):
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # -- Default paragraph font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # -- Heading styles --
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)  # BKA primary blue
        if level == 1:
            h.font.size = Pt(26)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(4)
        elif level == 2:
            h.font.size = Pt(16)
            h.font.italic = False
            h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        elif level == 3:
            h.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # -- Title Page --
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("YouTube Upload Metadata")
    run.font.name = "Calibri Light"
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
    run.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run(video_filename)
    run.font.name = "Calibri Light"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("BK Advisors")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)

    doc.add_page_break()

    # -- Video Title --
    doc.add_heading("Video Title", level=2)
    p = doc.add_paragraph(metadata["title"])
    p.paragraph_format.space_after = Pt(8)

    # -- Category --
    doc.add_heading("Category", level=2)
    p = doc.add_paragraph(metadata["category"])
    p.paragraph_format.space_after = Pt(8)

    # -- Description --
    doc.add_heading("Description", level=2)
    desc_paragraphs = metadata["description"].split("\n")
    for line in desc_paragraphs:
        if line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    # -- Tags --
    doc.add_heading("Tags", level=2)
    p = doc.add_paragraph(metadata["tags"])
    p.paragraph_format.space_after = Pt(8)

    # -- Video Chapters --
    doc.add_heading("Video Chapters", level=2)
    chapter_lines = metadata["chapters"].split("\n")
    for line in chapter_lines:
        if line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    # -- Thumbnail Suggestions --
    doc.add_heading("Thumbnail Suggestions", level=2)
    thumb_paragraphs = metadata["thumbnail_suggestions"].split("\n")
    for line in thumb_paragraphs:
        if line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    # -- Save --
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    base_dir = os.path.dirname(__file__)

    build_document(
        intro_video,
        "Intro - Storytelling with Data.mp4",
        os.path.join(base_dir, "youtube-metadata-intro-video.docx"),
    )

    build_document(
        full_lecture,
        "Full Lecture - Storytelling with Data.mp4",
        os.path.join(base_dir, "youtube-metadata-full-lecture.docx"),
    )
