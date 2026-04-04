"""
Generate Pre-Class Essay for "Storytelling with Data"
Output: pre-class-essay.docx
Style: Morgan Housel (conversational, story-driven, counterintuitive insights)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# -- Essay Content -----------------------------------------------------------

essay = {
    "title": "The Chart That Changed My Mind",
    "subtitle": "Why the best data stories don't start with data",
    "body": r"""
Last year I came across an article on Our World in Data — "Measles vaccines save millions of lives each year" by Saloni Dattani and Fiona Spooner. It had everything you'd want in a data story: clear writing, strong evidence, and a headline number that's hard to forget — 94 million lives saved over 50 years of vaccination programs.

But the thing that stopped me wasn't the number. It was a heatmap.

The article included a heatmap of measles cases across US states from 1929 to 2022. Before the vaccine, the map was a wall of color — cases everywhere, every year, in every state. Then, around 1963 when the vaccine was introduced, the color dropped off. A cliff. The before-and-after was so stark it didn't need a caption. You could see the intervention working.

I stared at that chart for a while. Then I thought: what does this look like for Africa?

Africa is where the measles burden is heaviest today, and it's where vaccination campaigns have saved the most lives — 29 million, according to the article. But I'd never seen the African story told with that same visual clarity. So I opened R, pulled the WHO data, and started building my own heatmap — forty-three years of data, every African country, side by side.

What came out of that afternoon became a project I'm still proud of. And working on it taught me something I keep re-learning: the most important part of a data story is the part that isn't data. It's the question that got you to open the dataset in the first place.

Most of us who work with data were trained in a linear process. Define the problem. Collect data. Clean it. Explore it. Visualize it. Report. It's a sensible pipeline, and it produces sensible outputs — dashboards, reports, slide decks full of charts. But it skips the question that makes the difference between a chart someone glances at and a chart someone can't stop thinking about: why should anyone care?

The heatmap worked not because of the data behind it or the tool I used to make it. It worked because I'd done three things before I ever opened a dataset.

First, I had a question — not a topic, a question. Not "measles in Africa" but "did the 2001 initiative actually cause the decline, or was it something else?" That's a question with stakes. It implies that the answer could go either way. It makes you lean in.

Second, I had a diagnosis. I didn't just show that measles declined. I showed why. The temporal structure of the heatmap — forty-three years, country by country — made the 2001 inflection point visible as a cliff edge. You could see the before and after. You could see which countries responded and which didn't. The chart wasn't just describing a trend. It was identifying a mechanism.

Third, I had an implication. The unspoken question hanging over the chart was: what happens if we let coverage slip? The heatmap answered it without saying a word. The countries that lagged were right there, in red, on the right side of the chart. The prediction was embedded in the pattern.

This is what I've come to think of as the full loop of data storytelling. Most data stories cover the middle — here's what the data shows. The good ones also cover the beginning — here's why it happened. The best ones complete the loop — here's what happens next.

It maps onto something I remember from my physics education. In physics, you don't just observe that a ball fell. You identify the force (diagnosis), describe the trajectory (description), and predict where it will land (prediction). The same structure applies to telling stories with data. A chart that only describes is like a physics experiment that only records the fall without understanding gravity.

The uncomfortable truth is that most of us spend our careers in the descriptive zone. We get very good at making charts that show what happened. We learn which chart types work for which data structures. We develop aesthetic taste — clean axes, readable labels, appropriate color palettes. All of this is valuable. But it's incomplete.

The diagnostic layer is where insight lives. When you can point to a chart and say not just "measles declined" but "measles declined because of this specific intervention, and here's how you can see the mechanism in the data" — that's when your audience stops nodding politely and starts taking notes.

The predictive layer is where action lives. When you can say "and if we don't maintain coverage, here's what the data suggests will happen" — that's when your audience stops taking notes and starts making phone calls.

I didn't learn this in a course. I learned it by making the mistake a hundred times — building beautiful charts that nobody acted on — and then watching someone with a simpler chart and a clearer story change a room.

Morgan Housel, one of my favorite writers, barely uses charts at all. His book "The Psychology of Money" has exactly zero data visualizations. Yet it's one of the most effective pieces of financial data storytelling ever written. He achieves with narrative what most of us try to achieve with ggplot2: he makes you feel the weight of a number. He does it with context ("This is how much a dollar in 1960 is worth today"), comparison ("Warren Buffett earned 97% of his wealth after age 65"), and empathy ("You're not irrational — you're just shaped by a different experience").

The lesson isn't that charts don't matter. They do. The lesson is that charts are amplifiers. They amplify a clear finding. They amplify a strong diagnosis. They amplify a well-grounded prediction. But they can't create any of those things. A chart without a story is decoration. A story without a chart can still change minds.

So where do you start? You start small. You start with one question that genuinely itches. Not the question your boss asked or the one in the terms of reference — the one that keeps you curious. You make it measurable. You look at the data with honest eyes, willing to be surprised. You declare what you found in one sentence — the "so what?" sentence. And then, when you're ready, you ask the two questions most data storytellers forget: Why did this happen? And what happens next?

The first chart I made in 2015 was a simple bar chart. It wasn't interactive. It didn't scroll. It wasn't built with a fancy tool. But it answered a specific question for a specific person, and it changed how that person thought about the problem.

That's data storytelling. Not the capital-S, ten-person-team kind. The kind you do with one question and one dataset and one person you're trying to reach. And it compounds. Each small project builds your eye for diagnosis, your instinct for prediction, your sense for what your reader actually needs. A decade later, the bar charts have become heatmaps and scrollytelling experiences. But the core hasn't changed: start with curiosity, end with an honest answer, and think about the person on the other side of the screen.
""",
}

# -- Document Builder --------------------------------------------------------


def build_document(essay, output_path):
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # -- Default paragraph font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # -- Heading styles --
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)  # BKA primary blue
        if level == 1:
            h.font.size = Pt(26)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(4)
        elif level == 2:
            h.font.size = Pt(16)
            h.font.italic = True
            h.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(18)
        elif level == 3:
            h.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # -- Title Page --
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Storytelling with Data")
    run.font.name = "Calibri Light"
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
    run.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("Pre-Class Essay")
    run.font.name = "Calibri Light"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("BK Advisors")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)

    doc.add_page_break()

    # -- Essay --
    doc.add_heading(essay["title"], level=1)
    doc.add_heading(essay["subtitle"], level=2)

    # Body text — split into paragraphs
    body = essay["body"].strip()
    paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]

    for para_text in paragraphs:
        clean_text = " ".join(para_text.split("\n"))
        p = doc.add_paragraph(clean_text)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)

    # -- Save --
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    out_dir = os.path.join(os.path.dirname(__file__), "..", "generated")
    os.makedirs(out_dir, exist_ok=True)
    output = os.path.join(out_dir, "pre-class-essay.docx")
    build_document(essay, output)
