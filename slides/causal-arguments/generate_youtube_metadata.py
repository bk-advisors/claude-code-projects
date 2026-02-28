"""
Generate YouTube Upload Metadata for Causal Arguments Videos
Output: youtube-metadata-intro-video.docx, youtube-metadata-full-lecture.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# -- Metadata Content -------------------------------------------------------

intro_video = {
    "title": "Causal Arguments for Consultants: Prove It, Don't Just Claim It | Lecture 2 Intro",
    "description": (
        "Why do some program evaluations get funded while others collect dust? "
        "The difference isn't the data — it's the argument.\n\n"
        "In this 5-minute introduction, Matthew Kuch shares the story of two competing reports "
        "from Ethiopia that landed on the same donor's desk — and explains why the team with "
        "weaker results got funded while the team with stronger data got rejected.\n\n"
        "This intro previews the full lecture on Making Causal Arguments That Hold Up, "
        "part of BK Advisors' professional skills series for global health consultants.\n\n"
        "YOU'LL LEARN:\n"
        "- The critical difference between correlation and causation\n"
        "- A 3-part diagnostic checklist (common cause, coincidence, reversed causation)\n"
        "- The two-part structure for building a causal argument\n"
        "- The removal test for confirming causation\n"
        "- A language calibration tool to match claim strength to evidence\n\n"
        "SERIES: Professional Skills for Global Health Consultants\n"
        "LECTURE 1: The Anatomy of a Good Point\n"
        "LECTURE 2: Making Causal Arguments That Hold Up (this video)\n"
        "COMING NEXT: Inference and Intervention (10-chapter course)\n\n"
        "ABOUT THE SPEAKER:\n"
        "Matthew Kuch has spent over a decade in global child health — at CHAI in Uganda, "
        "Gavi in Geneva managing country portfolios across Kenya and Lesotho, and at Deloitte "
        "managing grants for USAID and DFID. He has reviewed hundreds of program evaluations "
        "and seen firsthand which causal claims survive scrutiny and which don't.\n\n"
        "BK Advisors | Professional Development for Global Health Consultants"
    ),
    "tags": (
        "causal arguments, correlation vs causation, global health consulting, "
        "program evaluation, maternal and newborn health, MNH, Ethiopia, "
        "evidence-based consulting, causal reasoning, consulting skills, "
        "donor reports, health policy, professional development, "
        "BK Advisors, Matthew Kuch, CHEW, community health workers, "
        "causal inference, logical reasoning, argument structure"
    ),
    "category": "Education",
    "thumbnail_suggestions": (
        "TEXT OVERLAY: \"Correlation ≠ Causation\" in large white bold text\n\n"
        "VISUAL: Split-screen concept — left side shows a simple trend line "
        "(labeled \"Correlation\"), right side shows a chain of linked arrows "
        "(labeled \"Causation\"). Use BK Advisors primary blue (#005CB9) background.\n\n"
        "ALTERNATIVE: Photo of Matthew Kuch with text overlay \"Can You PROVE It?\" "
        "and a subtle \"≠\" symbol between two data charts."
    ),
    "chapters": (
        "0:00 — The Report That Almost Changed Everything\n"
        "0:30 — About the Speaker\n"
        "1:00 — What This Session Covers\n"
        "1:45 — The Framework: Three Alternatives to Rule Out\n"
        "3:45 — A Complete Causal Argument in 60 Seconds\n"
        "4:30 — Who This Is For\n"
        "5:00 — What You'll Walk Away With\n"
        "5:30 — Your Challenge"
    ),
}

full_lecture = {
    "title": "Making Causal Arguments That Hold Up | Full Lecture 2 — BK Advisors",
    "description": (
        "Every time you say \"X caused Y\" in a report, a presentation, or a policy recommendation, "
        "you're making a causal claim. This lecture teaches you how to make those claims hold up "
        "under scrutiny.\n\n"
        "Using real examples from maternal and newborn health consulting in Ethiopia, "
        "Matthew Kuch walks through a complete framework for building, testing, and defending "
        "causal arguments — the skill that separates funded programs from rejected proposals.\n\n"
        "WHAT YOU'LL LEARN:\n"
        "- The difference between correlation and causation (and why it matters for your career)\n"
        "- The Timeline Trap: why \"X happened before Y\" is never enough\n"
        "- Three alternatives to rule out before claiming causation:\n"
        "    • Common Cause — could a third factor drive both?\n"
        "    • Coincidence — could X and Y be unrelated?\n"
        "    • Reversed Causation — could the direction be backwards?\n"
        "- How to build a two-part causal argument (correlation + causal story)\n"
        "- The Removal Test: the strongest proof observational data can offer\n"
        "- Language calibration: matching your claim strength to your evidence\n"
        "- Five common mistakes that destroy credibility\n"
        "- A complete worked example you can model in your own reports\n\n"
        "SERIES: Professional Skills for Global Health Consultants\n"
        "PREREQUISITE: Lecture 1 — The Anatomy of a Good Point\n"
        "THIS VIDEO: Lecture 2 — Making Causal Arguments That Hold Up\n"
        "COMING NEXT: Inference and Intervention (10-chapter course on causal models)\n\n"
        "ABOUT THE SPEAKER:\n"
        "Matthew Kuch has spent over a decade in global child health — at CHAI in Uganda, "
        "Gavi in Geneva managing country portfolios across Kenya and Lesotho, and at Deloitte "
        "managing grants for USAID and DFID. He has reviewed hundreds of program evaluations "
        "and seen firsthand which causal claims survive scrutiny and which don't.\n\n"
        "BK Advisors | Professional Development for Global Health Consultants"
    ),
    "tags": (
        "causal arguments, correlation vs causation, global health consulting, "
        "program evaluation, maternal and newborn health, MNH, Ethiopia, "
        "evidence-based consulting, causal reasoning, consulting skills, "
        "donor reports, health policy, professional development, "
        "BK Advisors, Matthew Kuch, CHEW, community health workers, "
        "causal inference, logical reasoning, argument structure, "
        "removal test, common cause, coincidence, reversed causation, "
        "causal story, language calibration, program evaluation mistakes, "
        "health systems, woreda, HSTP, consulting framework"
    ),
    "category": "Education",
    "thumbnail_suggestions": (
        "TEXT OVERLAY: \"Making Causal Arguments That Hold Up\" in white bold text, "
        "with \"FULL LECTURE\" badge in BK Advisors green (#83BD00)\n\n"
        "VISUAL: The three-alternatives diagnostic framework from the whiteboard diagram — "
        "central question \"X caused Y?\" surrounded by three labeled boxes "
        "(Common Cause, Coincidence, Reversed Causation). "
        "Use BK Advisors primary blue (#005CB9) background.\n\n"
        "ALTERNATIVE: Matthew Kuch presenting with the framework visible on screen behind him, "
        "with text overlay \"Correlation ≠ Causation\" and the BK Advisors logo."
    ),
    "chapters": (
        "NOTE: Timestamps below are approximate based on lecture structure.\n"
        "Adjust to match the actual video after reviewing.\n\n"
        "0:00 — Introduction: Where We Left Off\n"
        "XX:XX — Why Causal Arguments Matter\n"
        "XX:XX — The Core Problem: Correlation vs. Causation\n"
        "XX:XX — The Timeline Trap\n"
        "XX:XX — Three Alternatives to Rule Out\n"
        "XX:XX — Alternative 1: Common Cause\n"
        "XX:XX — Alternative 2: Coincidence\n"
        "XX:XX — Alternative 3: Reversed Causation\n"
        "XX:XX — The Diagnostic Checklist\n"
        "XX:XX — Building a Causal Argument: The Two-Part Structure\n"
        "XX:XX — What Makes a Good Causal Story\n"
        "XX:XX — The Removal Test\n"
        "XX:XX — Putting It All Together: Complete Example\n"
        "XX:XX — Five Common Mistakes\n"
        "XX:XX — Calibrate Your Language\n"
        "XX:XX — Key Takeaways\n"
        "XX:XX — What's Next: Inference and Intervention"
    ),
}

# -- Document Builder -------------------------------------------------------


def build_document(metadata, video_filename, output_path):
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # -- Default paragraph font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # -- Heading styles --
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)  # BKA primary blue
        if level == 1:
            h.font.size = Pt(26)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(4)
        elif level == 2:
            h.font.size = Pt(16)
            h.font.italic = False
            h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        elif level == 3:
            h.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # -- Title Page --
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("YouTube Upload Metadata")
    run.font.name = "Calibri Light"
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
    run.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run(video_filename)
    run.font.name = "Calibri Light"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("BK Advisors")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)

    doc.add_page_break()

    # -- Video Title --
    doc.add_heading("Video Title", level=2)
    p = doc.add_paragraph(metadata["title"])
    p.paragraph_format.space_after = Pt(8)

    # -- Category --
    doc.add_heading("Category", level=2)
    p = doc.add_paragraph(metadata["category"])
    p.paragraph_format.space_after = Pt(8)

    # -- Description --
    doc.add_heading("Description", level=2)
    desc_paragraphs = metadata["description"].split("\n")
    for line in desc_paragraphs:
        if line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    # -- Tags --
    doc.add_heading("Tags", level=2)
    p = doc.add_paragraph(metadata["tags"])
    p.paragraph_format.space_after = Pt(8)

    # -- Video Chapters --
    doc.add_heading("Video Chapters", level=2)
    chapter_lines = metadata["chapters"].split("\n")
    for line in chapter_lines:
        if line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    # -- Thumbnail Suggestions --
    doc.add_heading("Thumbnail Suggestions", level=2)
    thumb_paragraphs = metadata["thumbnail_suggestions"].split("\n")
    for line in thumb_paragraphs:
        if line.strip():
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

    # -- Save --
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    base_dir = os.path.dirname(__file__)

    build_document(
        intro_video,
        "Intro to Lecture 2 - Causal Arguments.mp4",
        os.path.join(base_dir, "youtube-metadata-intro-video.docx"),
    )

    build_document(
        full_lecture,
        "Full Lecture 2 - Causal Arguments.mp4",
        os.path.join(base_dir, "youtube-metadata-full-lecture.docx"),
    )
