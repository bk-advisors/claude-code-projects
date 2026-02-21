"""
Generate Pre-Class Essay for "Making Causal Arguments That Hold Up"
Output: pre-class-essay.docx
Style: Morgan Housel (conversational, story-driven, counterintuitive insights)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# -- Essay Content ----------------------------------------------------------

essay = {
    "title": "The Report That Almost Changed Everything",
    "subtitle": "Why proving causation is harder \u2014 and more important \u2014 than you think",
    "body": r"""
Two reports landed on the same desk in Addis Ababa within a week of each other.

The first was impressive. Fifty pages. A program manager in Ethiopia had deployed Community Health Extension Workers across twelve woredas in Amhara region. Eighteen months in, the numbers looked extraordinary: maternal mortality had dropped 28% in the program woredas. The report documented every data point, every trend line, every before-and-after comparison. The conclusion was bold: "Our CHEW deployment in Ethiopia caused a significant reduction in maternal mortality."

The donor sent it to their technical advisor. The advisor sent back one paragraph: "You've shown a correlation. You haven't shown causation. The government of Ethiopia also rolled out a new emergency obstetric care protocol in the same woredas during the same period. How do you know it was your CHEWs and not the protocol?"

The program manager was stunned. The data was right there. Twenty-eight percent. Twelve woredas. Eighteen months. How could anyone argue with that?

But the advisor wasn't arguing with the data. She was arguing with the logic. Two things had happened at the same time in the same places in Ethiopia, and the report had simply assumed that one of them caused the improvement. That's not proof. That's a coincidence dressed up as a conclusion.

The second report arrived a few days later from a different team working in Ethiopia's SNNPR. Their results were more modest \u2014 a 15% reduction in neonatal mortality. But the report was eight pages, not fifty. And it did something the first report didn't: it made an argument, not just a claim.

It showed the correlation. It traced the causal mechanism step by step: CHEWs in Ethiopia conducted home visits, identified danger signs early, referred mothers to health centers faster, reduced the first and second delays in emergency obstetric care, which reduced hemorrhage deaths. Each link in the chain was plausible and supported.

Then it addressed the obvious counterarguments. Could the improvement have a common cause? They checked: control woredas in Ethiopia with similar baseline resources showed no improvement. Could it be coincidence? A 15% drop was statistically significant and beyond normal quarterly variation. Could the causation be reversed? The woredas were randomly assigned, not self-selected.

And then the clincher: when CHEWs were temporarily reassigned from two facilities in Ethiopia, outcomes deteriorated within weeks. When they returned, outcomes improved again. Remove the cause, the effect disappears. Restore the cause, the effect returns.

The donor funded the scale-up within a month.

Here is the counterintuitive lesson: the team with weaker results got funded, and the team with stronger results got rejected. The difference was not the data. The difference was the argument.

This happens more often than anyone in global health likes to admit. We are trained to collect data, run analyses, and present findings. We are not trained to make causal arguments. And there is a chasm between showing that two things happened together and proving that one caused the other.

The chasm has a name: the correlation-causation gap. And it has three trapdoors.

The first is common cause. Two things can move together because a third thing is driving both. In Ethiopia, woredas with more NGO programs tend to have better health outcomes. But both might be caused by government prioritization under Ethiopia's Health Sector Transformation Plan. The NGOs didn't cause the outcomes. Ethiopia's government attention caused both the NGO presence and the outcomes. If you don't look for this, you won't find it, and your claim will be wrong.

The second is coincidence. Especially in small samples \u2014 one health center in Ethiopia, one quarter of data \u2014 random variation can masquerade as a pattern. A health center in Ethiopia introduced a new patient register the same month neonatal mortality dropped. No mechanism connects a register format to newborn survival. But if you're not careful, you'll write that up as a finding.

The third is reversed causation. Health posts in Ethiopia with more supervision visits have better outcomes. Obviously supervision improves performance, right? Maybe. Or maybe supervisors visit high-performing health posts more often because they're easier to reach and more cooperative. The causation runs backwards.

These three trapdoors are not obscure academic concerns. They are the reason good programs in Ethiopia lose funding and bad arguments win the room. If you cannot rule them out, your causal claim is standing on sand.

So what does a good causal argument look like? It has two parts.

First, show the correlation. X and Y move together. Use your data. This is necessary but not sufficient \u2014 it's the starting point, not the finish line.

Second, tell the causal story. Trace the mechanism step by step. How does X produce Y? What is the chain of events? Each link should be plausible. Each link should, ideally, be testable.

And then apply what you might call the removal test: if you take X away, does Y disappear? This is the closest thing to proof that observational data can offer. It's what the second team in Ethiopia did when they temporarily reassigned their CHEWs.

There is one more thing the second team did that the first team didn't: they calibrated their language. They didn't say "caused." They said "likely contributed to." That precision matters. Stakeholders in Ethiopia and beyond respect honesty about uncertainty. Overclaiming destroys credibility faster than any methodological flaw.

The program manager who wrote the first report eventually rewrote it. She kept the same data but restructured the argument. She traced the causal mechanism. She addressed the alternative explanations. She softened her language from "caused" to "likely contributed to" and explained why. The revised version was funded six months later.

Same data. Different argument. Different outcome.

The lesson is not that data doesn't matter. It matters enormously. The lesson is that data without a causal argument is just a pile of numbers, and a causal argument without a structure is just an assertion. The teams that change things in Ethiopia \u2014 the ones that get funded, scaled, and taken seriously \u2014 are the ones that can show the pattern, explain the mechanism, and rule out the alternatives.

A causal argument that holds up does three things: shows the pattern, explains the mechanism, and rules out the alternatives.

The next time you write a report, a recommendation, or a slide deck, don't just show the data. Make the argument. Trace the chain. Address the counterarguments. And calibrate your language. Because in the end, the difference between a report that changes policy and a report that collects dust is not the strength of the data \u2014 it's the strength of the reasoning.
""",
}

# -- Document Builder -------------------------------------------------------


def build_document(essay, output_path):
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # -- Default paragraph font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # -- Heading styles --
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)  # BKA primary blue
        if level == 1:
            h.font.size = Pt(26)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(4)
        elif level == 2:
            h.font.size = Pt(16)
            h.font.italic = True
            h.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(18)
        elif level == 3:
            h.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # -- Title Page --
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Making Causal Arguments That Hold Up")
    run.font.name = "Calibri Light"
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
    run.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("Pre-Class Essay")
    run.font.name = "Calibri Light"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("BK Advisors")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)

    doc.add_page_break()

    # -- Essay --
    doc.add_heading(essay["title"], level=1)
    doc.add_heading(essay["subtitle"], level=2)

    # Body text -- split into paragraphs
    body = essay["body"].strip()
    paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]

    for para_text in paragraphs:
        # Handle line continuations within paragraphs
        clean_text = " ".join(para_text.split("\n"))
        p = doc.add_paragraph(clean_text)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)

    # -- Save --
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    output = os.path.join(os.path.dirname(__file__), "pre-class-essay.docx")
    build_document(essay, output)
